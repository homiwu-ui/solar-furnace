#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
追日電路設計說明 Word 檔
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '微軟正黑體'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_h1(doc, text, color=RGBColor(0, 102, 153)):
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.color.rgb = color

def add_h2(doc, text, color=RGBColor(0, 102, 51)):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.color.rgb = color

def add_h3(doc, text, color=RGBColor(51, 102, 153)):
    h = doc.add_heading(text, level=3)
    for run in h.runs:
        run.font.color.rgb = color

def add_box(doc, title, text, bg='E8F4FD', tc='003366'):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, bg)
    p = cell.paragraphs[0]
    run = p.add_run(title + '\n')
    run.bold = True
    run.font.color.rgb = RGBColor(int(tc[:2],16), int(tc[2:4],16), int(tc[4:],16))
    run.font.size = Pt(13)
    run2 = p.add_run(text)
    run2.font.size = Pt(11)
    run2.font.color.rgb = RGBColor(51, 51, 51)
    doc.add_paragraph('')

def add_parts_table(doc, items):
    table = doc.add_table(rows=1+len(items), cols=4)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['物品', '規格', '數量', '備註']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, '1F4E79')
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(11)
    for ri, rd in enumerate(items):
        for ci, val in enumerate(rd):
            table.rows[ri+1].cells[ci].text = val
            for p in table.rows[ri+1].cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph('')

def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = '新細明體'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(51, 51, 51)
    return p

# ============================================================
# 封面
# ============================================================
for _ in range(6):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('追日電路設計說明')
r.bold = True
r.font.size = Pt(36)
r.font.color.rgb = RGBColor(0, 102, 153)

doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('用光和電讓太陽爐自動追蹤太陽')
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(102, 102, 102)
r.italic = True

doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('本文件包含：電路原理、材料清單、接線圖、組裝步驟')
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(102, 102, 102)

doc.add_page_break()

# ============================================================
# 目錄
# ============================================================
add_h1(doc, '本文件目錄')
sections = [
    ('一', '追日電路原理', '為什麼光敏電阻可以讓太陽爐自動轉？'),
    ('二', '材料清單', '需要哪些零件？去哪裡買？'),
    ('三', '核心電路圖', '雙運放比較器電路'),
    ('四', '電源管理電路', '太陽能板 + 穩壓 + 電池'),
    ('五', '馬達驅動電路', 'L293D 驅動馬達正反轉'),
    ('六', '完整接線圖', '所有零件連在一起的總圖'),
    ('七', '組裝步驟', '一步一步教你接'),
    ('八', '測試與除錯', '怎麼知道有沒有成功？'),
]
for num, title, desc in sections:
    p = doc.add_paragraph()
    r = p.add_run(f'{num}、{title}')
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0, 102, 153)
    doc.add_paragraph(f'    {desc}')

doc.add_page_break()

# ============================================================
# 一、追日電路原理
# ============================================================
add_h1(doc, '一、追日電路原理')

add_h2(doc, '核心概念：用光來「指路」')
doc.add_paragraph(
    '太陽從東邊升起，慢慢移到西邊。如果太陽能板固定不動，'
    '一天當中只有正午能完全對準太陽，效率很低。'
    '「追日系統」就是讓太陽能板自動跟著太陽轉，永遠對準太陽，'
    '這樣發電效率可以提升 30% 以上！'
)

add_box(doc, '追日的秘密',
    '祕密在「兩顆光敏電阻」！\n'
    '把它們放在太陽能板的左邊和右邊，\n'
    '當太陽在左邊時，左邊的光敏電阻感測到更強的光，\n'
    '電路就會自動驅動馬達把太陽能板轉向左邊。\n'
    '就這樣一直追著太陽跑！',
    'E8F4FD', '003366')

add_h2(doc, '運作原理圖')

code_block(doc,
    '  電路運作流程：\n\n'
    '    ☀️ 太陽在左邊          ☀️ 太陽在正中央         ☀️ 太陽在右邊\n'
    '        ↓                      ↓                      ↓\n'
    '    LDR1 亮                 LDR1 ≈ LDR2            LDR2 亮\n'
    '    LDR2 暗                 兩邊一樣亮              LDR1 暗\n'
    '        ↓                      ↓                      ↓\n'
    '    電壓差 → 馬達轉        電壓差 ≈ 0              電壓差 → 馬達轉\n'
    '    (向左轉)               馬達停止                 (向右轉)\n'
    '    (追蹤太陽)             (已對準！)               (追蹤太陽)'
)

add_h2(doc, '關鍵零件介紹')

add_h3(doc, '光敏電阻（LDR）')
doc.add_paragraph(
    '光敏電阻是一種「光越強，電阻越小」的零件。\n'
    '  - 沒光時：電阻很大（約 1MΩ），幾乎不導電\n'
    '  - 強光時：電阻很小（約 5KΩ），導電性好\n'
    '我們利用這個特性，把光的強弱變成電壓的高低。'
)

add_h3(doc, '運算放大器 / 比較器（LM358 / LM393）')
doc.add_paragraph(
    '比較器的工作很簡單：比較兩個電壓，決定輸出高或低。\n'
    '  - 如果「正端電壓 > 負端電壓」→ 輸出高電位\n'
    '  - 如果「正端電壓 < 負端電壓」→ 輸出低電位\n'
    '我們把兩顆光敏電阻的電壓分別接到比較器的兩端，'
    '它就會自動判斷太陽在哪一邊！'
)

add_h3(doc, '減速馬達')
doc.add_paragraph(
    '太陽每分鐘只移動 0.25 度，所以馬達轉速要很慢。\n'
    '「減速馬達」裡面有齒輪組，可以把高轉速變成低轉速但力氣大。\n'
    '推薦使用 TT 減速馬達（1:48），轉速約 4RPM，力氣足夠推動太陽爐。'
)

doc.add_page_break()

# ============================================================
# 二、材料清單
# ============================================================
add_h1(doc, '二、材料清單')

add_h2(doc, 'A. 感測元件')
add_parts_table(doc, [
    ['光敏電阻（LDR）', '5516 或 GL5528', '2顆', '亮阻 5-10KΩ，暗阻 1MΩ'],
    ['固定電阻', '10KΩ 1/4W', '4顆', '分壓用，搭配 LDR'],
    ['固定電阻', '330Ω 1/4W', '2顆', 'LED 限流用'],
    ['LED 指示燈', '紅色 3mm', '1顆', '電源指示'],
])

add_h2(doc, 'B. 控制電路')
add_parts_table(doc, [
    ['比較器 IC', 'LM393（雙比較器）', '1顆', '也可用 LM358 運放'],
    ['IC 底座', 'DIP-8', '1個', '方便更換 IC'],
    ['穩壓 IC', 'LM7805（5V 輸出）', '1顆', 'TO-220 封裝'],
    ['電容', '100μF 16V', '2顆', '電源濾波用'],
    ['電容', '0.1μF（104）', '2顆', '高頻濾波用'],
])

add_h2(doc, 'C. 電源系統')
add_parts_table(doc, [
    ['太陽能板', '6V 3W 以上', '1片', '供電給系統 + 充電'],
    ['鋰電池', '18650 3.7V 2000mAh', '1-2顆', '備援電源'],
    ['電池座', '18650 單顆座', '1個', '含導線'],
    ['二極體', '1N4007', '1顆', '防逆流保護'],
    ['電源開關', '小型撥動開關', '1個', '系統總開關'],
])

add_h2(doc, 'D. 動力系統')
add_parts_table(doc, [
    ['減速馬達', 'TT 6V 1:48', '1-2顆', '水平 + 垂直各一顆'],
    ['馬達驅動 IC', 'L293D', '1顆', 'DIP-16 封裝，可正反轉'],
    ['IC 底座', 'DIP-16', '1個', '方便更換 IC'],
    ['杜邦線', '公對公 / 公對母', '各 10條', '連接電路'],
    ['麵包板', '830 孔', '1個', '測試用，正式焊接用萬用板'],
])

add_h2(doc, 'E. 購買管道')

table = doc.add_table(rows=5, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['管道', '適合買什麼', '備註']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    set_cell_shading(cell, '1F4E79')
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(11)

data = [
    ['蝦皮購物', 'LDR、LM393、馬達、太陽能板', '最便宜，搜尋「光敏電阻」「TT 減速馬達」'],
    ['全國電子 / 3C好室', '焊接工具、電阻、電容', '實體店面，馬上拿到'],
    ['光華商場', 'IC 晶片、萬用板、電子零件', '台北的話直接去逛'],
    ['淘寶', 'Arduino 套件、追光模組', '品項最多，但要等寄送'],
]
for ri, rd in enumerate(data):
    for ci, val in enumerate(rd):
        table.rows[ri+1].cells[ci].text = val
        for p in table.rows[ri+1].cells[ci].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
doc.add_paragraph('')

doc.add_page_break()

# ============================================================
# 三、核心電路圖
# ============================================================
add_h1(doc, '三、核心電路圖')

add_h2(doc, '3.1 雙運放比較器電路（最簡版）')

doc.add_paragraph(
    '這是最基本的追日電路，只需要 2 顆光敏電阻、2 個電阻、1 顆 LM358，'
    '就能讓馬達自動轉動。'
)

code_block(doc,
    '  ===== 雙運放比較器電路 =====\n\n'
    '        VCC (+5V)\n'
    '         │\n'
    '    ┌────┴────┐\n'
    '    │         │\n'
    '   [R1]     [R2]\n'
    '   10KΩ     10KΩ\n'
    '    │         │\n'
    '    ├────┬────┤\n'
    '    │    │    │\n'
    '  [LDR1] │  [LDR2]\n'
    '    │    │    │\n'
    '    ▼    │    ▼\n'
    '   GND   │   GND\n'
    '         │\n'
    '    ┌────┴────────────────────┐\n'
    '    │        LM358           │\n'
    '    │  ┌──────────┐          │\n'
    '    │  │  運放 A   │          │\n'
    '    │  │ IN+  OUT├─────┐    │\n'
    '    │  │ IN-     │     │    │\n'
    '    │  └──────────┘     │    │\n'
    '    │  ┌──────────┐     │    │\n'
    '    │  │  運放 B   │     │    │\n'
    '    │  │ IN+     │     │    │\n'
    '    │  │ IN-  OUT├─────┤    │\n'
    '    │  └──────────┘     │    │\n'
    '    │                   │    │\n'
    '    │  OUT A ───────────┤    │\n'
    '    │  OUT B ───────────┤    │\n'
    '    └───────────────────┼────┘\n'
    '                        │\n'
    '                      ┌─┴─┐\n'
    '                      │ DC │\n'
    '                      │馬達│\n'
    '                      └───┘'
)

add_box(doc, '運作邏輯',
    'LDR1 亮（太陽在左）→ LDR1 阻值↓ → 分壓↓ → OUT_A > OUT_B → 馬達正轉\n'
    'LDR2 亮（太陽在右）→ LDR2 阻值↓ → 分壓↓ → OUT_B > OUT_A → 馬達反轉\n'
    '兩邊一樣亮 → OUT_A ≈ OUT_B → 馬達停止（已對準）',
    'D4EDDA', '155724')

doc.add_page_break()

# ============================================================
# 四、電源管理電路
# ============================================================
add_h1(doc, '四、電源管理電路')

add_h2(doc, '4.1 太陽能板 + 穩壓電路')

doc.add_paragraph(
    '太陽能板輸出的電壓會隨著光線強弱變化（3V~7V），'
    '但我們的電路需要穩定的 5V。所以需要「穩壓 IC」來穩定電壓。'
)

code_block(doc,
    '  ===== 電源管理電路 =====\n\n'
    '    ┌─────────────────────────────┐\n'
    '    │      太陽能板 6V/3W         │\n'
    '    │        ┌─────┐             │\n'
    '    │        │ ☀️  │             │\n'
    '    │        └──┬──┘             │\n'
    '    └───────────┼────────────────┘\n'
    '                │\n'
    '              [D1] 1N4007 ← 防逆流（防止電池電回灌到太陽能板）\n'
    '                │\n'
    '          ┌─────┴─────┐\n'
    '          │           │\n'
    '        [C1]        [U1]\n'
    '      100μF       LM7805\n'
    '      濾波電容     穩壓 IC\n'
    '          │           │\n'
    '          └─────┬─────┘\n'
    '                │\n'
    '             +5V bus\n'
    '                │\n'
    '          ┌─────┴─────┐\n'
    '          │           │\n'
    '        [C2]        [C3]\n'
    '       0.1μF      100μF\n'
    '        高頻濾波    輸出濾波\n'
    '          │           │\n'
    '         GND         GND'
)

add_h2(doc, '4.2 含電池備援的完整電源電路')

code_block(doc,
    '  ===== 完整電源電路（含 18650 電池備援）=====\n\n'
    '    ┌──────────────────────┐\n'
    '    │    太陽能板 6V/3W    │\n'
    '    └──────────┬───────────┘\n'
    '               │\n'
    '             [D1] 1N4007\n'
    '               │\n'
    '    ┌──────────┴──────────┐\n'
    '    │                     │\n'
    '  [R_charge]            [U1] LM7805\n'
    '   充電電阻               │\n'
    '   (可省略)            +5V bus\n'
    '    │                     │\n'
    '  ┌─┴─┐              ┌────┴────┐\n'
    '  │18650│             │  LM393  │\n'
    '  │電池 │             │  比較器  │\n'
    '  └───┘              └────┬────┘\n'
    '                          │\n'
    '                       [DC馬達]\n'
    '                          │\n'
    '                         GND\n\n'
    '  說明：\n'
    '  - 白天：太陽能板供電 + 給電池充電\n'
    '  - 晚上或陰天：電池備援供電\n'
    '  - D1 防止電池電流回灌到太陽能板\n'
    '  - LM7805 穩定輸出 5V 給控制電路'
)

add_box(doc, '重要提醒',
    'LM7805 的輸入電壓必須比輸出電壓高 2V 以上（即至少 7V），'
    '所以如果太陽能板只有 6V，穩壓 IC 可能無法正常工作。\n'
    '解決方案：\n'
    '  1. 用 12V 太陽能板 + LM7805\n'
    '  2. 改用 DC-DC 升壓模組（6V → 9V）再接 LM7805\n'
    '  3. 直接用 5V 降壓模組（如 LM2596）取代 LM7805',
    'FFF3CD', '856404')

doc.add_page_break()

# ============================================================
# 五、馬達驅動電路
# ============================================================
add_h1(doc, '五、馬達驅動電路')

add_h2(doc, '5.1 為什麼需要馬達驅動 IC？')
doc.add_paragraph(
    '比較器（LM393）的輸出電流很小（約 20mA），'
    '不足以直接驅動馬達（需要 100-200mA）。'
    '所以我們需要用「馬達驅動 IC」來放大電流，'
    '同時支援馬達正反轉。'
)

add_h2(doc, '5.2 L293D 馬達驅動電路')

code_block(doc,
    '  ===== L293D 接線圖 =====\n\n'
    '    L293D 腳位圖：\n'
    '    ┌─────────────────────────┐\n'
    '    │  ┌───┐                  │\n'
    '    │  │ 1 │ Enable1    VCC │16│──→ +5V (邏輯電源)\n'
    '    │  │ 2 │ Input1   Output4│15│\n'
    '    │  │ 3 │ Output1  Output4│14│\n'
    '    │  │ 4 │ GND          GND│13│──→ GND\n'
    '    │  │ 5 │ GND          GND│12│──→ GND\n'
    '    │  │ 6 │ Output2  Output3│11│\n'
    '    │  │ 7 │ Input2    Input3│10│\n'
    '    │  │ 8 │ VS       Enable2│ 9│──→ +6V (馬達電源)\n'
    '    │  └───┘                  │\n'
    '    └─────────────────────────┘\n\n'
    '  接線方式：\n'
    '  ┌────────────┐    ┌────────────┐    ┌────────────┐\n'
    '  │   LM393    │    │   L293D    │    │   DC 馬達  │\n'
    '  │            │    │            │    │            │\n'
    '  │ OUT_A ─────┼───→│ Input1 (2) │    │            │\n'
    '  │            │    │ Output1(3)─┼───→│ 紅線 (+)  │\n'
    '  │            │    │            │    │            │\n'
    '  │            │    │ Output2(6)─┼───→│ 黑線 (-)  │\n'
    '  │            │    │ Input2 (7) │    │            │\n'
    '  │            │    │   → GND   │    │            │\n'
    '  │            │    │            │    │            │\n'
    '  │            │    │ Enable1(1) │    │            │\n'
    '  │            │    │   → +5V   │    │            │\n'
    '  │            │    │            │    │            │\n'
    '  │            │    │ VS (8)     │    │            │\n'
    '  │            │    │   → +6V   │    │            │\n'
    '  └────────────┘    └────────────┘    └────────────┘'
)

add_h2(doc, '5.3 馬達正反轉邏輯')

table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['LM393 輸出', 'L293D 狀態', '馬達動作']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    set_cell_shading(cell, '1F4E79')
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(11)

data = [
    ['OUT_A > OUT_B', 'Input1=高, Input2=低', '正轉（向東/向上）'],
    ['OUT_A < OUT_B', 'Input1=低, Input2=高', '反轉（向西/向下）'],
    ['OUT_A ≈ OUT_B', 'Input1=低, Input2=低', '停止（已對準）'],
]
for ri, rd in enumerate(data):
    for ci, val in enumerate(rd):
        table.rows[ri+1].cells[ci].text = val
        for p in table.rows[ri+1].cells[ci].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
doc.add_paragraph('')

doc.add_page_break()

# ============================================================
# 六、完整接線圖
# ============================================================
add_h1(doc, '六、完整接線圖')

add_h2(doc, '6.1 全系統方塊圖')

code_block(doc,
    '  ===== 全系統方塊圖 =====\n\n'
    '    ┌───────────────────────────────────────────────┐\n'
    '    │                 太陽能板 6V                     │\n'
    '    └─────────────────────┬─────────────────────────┘\n'
    '                          │\n'
    '                        [D1]\n'
    '                          │\n'
    '                    ┌─────┴─────┐\n'
    '                    │  LM7805   │\n'
    '                    │  穩壓 5V   │\n'
    '                    └─────┬─────┘\n'
    '                          │\n'
    '    ┌─────────────────────┼─────────────────────┐\n'
    '    │                     │                     │\n'
    '    ▼                     ▼                     ▼\n'
    '  [C1]                  [C2]                 [C3]\n'
    ' 100μF                 0.1μF                100μF\n'
    '    │                     │                     │\n'
    '    └──────────┬──────────┴──────────┬──────────┘\n'
    '               │                     │\n'
    '          ┌────┴────┐           ┌────┴────┐\n'
    '          │  LM393  │           │  LM393  │\n'
    '          │ (水平)  │           │ (垂直)  │\n'
    '          └────┬────┘           └────┬────┘\n'
    '               │                     │\n'
    '    ┌──────────┴──────────┐ ┌────────┴─────────┐\n'
    '    │                     │ │                  │\n'
    '  [LDR1]              [LDR2] [LDR3]         [LDR4]\n'
    '  (左)                (右)   (上)           (下)\n'
    '    │                     │ │                  │\n'
    '    ▼                     ▼ ▼                  ▼\n'
    '   GND                   GND GND              GND\n\n'
    '  感測電路（4顆 LDR 感測上下左右）'
)

code_block(doc,
    '  ===== 馬達驅動部分 =====\n\n'
    '          ┌────┴────┐           ┌────┴────┐\n'
    '          │  LM393  │           │  LM393  │\n'
    '          │ (水平)  │           │ (垂直)  │\n'
    '          │  OUT_A  │           │  OUT_A  │\n'
    '          └────┬────┘           └────┬────┘\n'
    '               │                     │\n'
    '          ┌────▼────┐           ┌────▼────┐\n'
    '          │ L293D   │           │ L293D   │\n'
    '          │(水平驅動)│           │(垂直驅動)│\n'
    '          └────┬────┘           └────┬────┘\n'
    '               │                     │\n'
    '          ┌────▼────┐           ┌────▼────┐\n'
    '          │ 水平馬達 │           │ 垂直馬達 │\n'
    '          │ (東西向) │           │ (仰角)  │\n'
    '          └─────────┘           └─────────┘'
)

doc.add_page_break()

# ============================================================
# 七、組裝步驟
# ============================================================
add_h1(doc, '七、組裝步驟')

add_h2(doc, '7.1 準備工作')
steps_prep = [
    '確認所有零件都已購齊（參考第二章材料清單）',
    '準備工具：電烙鐵、焊錫、萬用表、麵包板、杜邦線',
    '在麵包板上先測試電路，確認正常後再焊接',
]
for s in steps_prep:
    doc.add_paragraph(s, style='List Number')

add_h2(doc, '7.2 在麵包板上搭建電路')

add_h3(doc, 'Step 1：搭建電源電路')
steps_1 = [
    '插入 LM7805 穩壓 IC（中間腳 GND，左邊腳 IN，右邊腳 OUT）',
    '在 IN 腳和 GND 之間接一個 100μF 電容（長腳接 IN，短腳接 GND）',
    '在 OUT 腳和 GND 之間接一個 100μF 電容',
    'IN 腳接太陽能板正極（經過 1N4007 二極體）',
    'OUT 腳就是穩定的 5V 輸出',
    '接上 LED 和 330Ω 電阻到 5V，作為電源指示燈',
]
for s in steps_1:
    doc.add_paragraph(s, style='List Number')

add_h3(doc, 'Step 2：搭建感測電路')
steps_2 = [
    '在麵包板上插 2 顆 10KΩ 電阻，一端接 5V，另一端接一行',
    '在每行的另一端各接一顆 LDR',
    'LDR 的另一端接 GND',
    '這樣就形成 2 組「5V → 電阻 → LDR → GND」的分壓電路',
    '用萬用表量測中間點的電壓，用手遮擋 LDR 應該能看到電壓變化',
]
for s in steps_2:
    doc.add_paragraph(s, style='List Number')

code_block(doc,
    '  麵包板接線示意：\n\n'
    '    5V ───┬───────────────┬────────\n'
    '           │               │\n'
    '         [R1] 10KΩ      [R2] 10KΩ\n'
    '           │               │\n'
    '    ┌──────┤          ┌────┤\n'
    '    │ LM393│          │LM393│\n'
    '    │ IN+  │          │IN+  │\n'
    '    └──────┘          └─────┘\n'
    '           │               │\n'
    '         [LDR1]         [LDR2]\n'
    '           │               │\n'
    '    GND ───┴───────────────┴────────'
)

add_h3(doc, 'Step 3：接上 LM393 比較器')
steps_3 = [
    '把 LM393 插入麵包板（注意方向，缺口朝左）',
    'Pin 8 接 5V，Pin 4 接 GND',
    'Pin 3 (IN+) 接 LDR1 和 R1 的中間點',
    'Pin 2 (IN-) 接 LDR2 和 R2 的中間點',
    'Pin 1 (OUT) 接到 L293D 的 Input1',
    '重複另一組：Pin 5 (IN+)、Pin 6 (IN-)、Pin 7 (OUT)',
]
for s in steps_3:
    doc.add_paragraph(s, style='List Number')

add_h3(doc, 'Step 4：接上 L293D 馬達驅動')
steps_4 = [
    '把 L293D 插入麵包板',
    'Pin 1 (Enable1) 接 5V',
    'Pin 16 (VCC) 接 5V',
    'Pin 8 (VS) 接電池正極（6V）',
    'Pin 4, 5, 12, 13 全部接 GND',
    'Pin 2 (Input1) 接 LM393 的 OUT_A',
    'Pin 7 (Input2) 接 GND（固定為低電位）',
    'Pin 3 (Output1) 接馬達紅線',
    'Pin 6 (Output2) 接馬達黑線',
]
for s in steps_4:
    doc.add_paragraph(s, style='List Number')

doc.add_page_break()

# ============================================================
# 八、測試與除錯
# ============================================================
add_h1(doc, '八、測試與除錯')

add_h2(doc, '8.1 測試步驟')

tests = [
    ('電源測試', '接上太陽能板或電池，LED 應該亮起。用萬用表量 LM7805 輸出應為 5V ±0.2V'),
    ('感測測試', '用手電筒照 LDR1，量 LM393 Pin 1 輸出應為高電位（~5V）；遮擋 LDR1 應變低電位（~0V）'),
    ('馬達測試', '用手電筒照左邊 LDR，馬達應正轉；照右邊 LDR，馬達應反轉'),
    ('追日測試', '把手電筒放在太陽能板左邊，馬達應自動轉向左邊追蹤光源'),
    ('停止測試', '把手電筒放在正中央，馬達應停止轉動'),
]
for name, desc in tests:
    p = doc.add_paragraph()
    r = p.add_run(name + '：')
    r.bold = True
    r.font.color.rgb = RGBColor(0, 102, 51)
    p.add_run(desc)

add_h2(doc, '8.2 常見問題與除錯')

problems = [
    ('LED 不亮', '電源沒接好、太陽能板沒光、電池沒電 → 檢查接線和電源'),
    ('LED 亮但馬達不動', 'L293D 沒接好、VS 腳沒接電源 → 檢查 L293D 接線'),
    ('馬達一直轉不停', 'LDR 沒裝好、電阻沒接 → 檢查分壓電路和 LDR 位置'),
    ('馬達轉反了', '馬達紅黑線接反 → 把 Output1 和 Output2 的線對調'),
    ('馬達轉太快', '沒有用減速馬達 → 換成 TT 減速馬達或 N20 馬達'),
    ('追蹤不靈敏', 'LDR 間距太近 → 把兩顆 LDR 拉開至少 5cm'),
]
for prob, sol in problems:
    p = doc.add_paragraph()
    r = p.add_run('Q：' + prob)
    r.bold = True
    r.font.color.rgb = RGBColor(0, 51, 153)
    doc.add_paragraph('A：' + sol)
    doc.add_paragraph('')

add_box(doc, '升級建議',
    '如果你已經完成基本追日功能，可以考慮：\n'
    '  1. 加入 Arduino 控制 → 更精準，可以寫程式調整\n'
    '  2. 加入窗口比較器 → 避免馬達一直微調耗電\n'
    '  3. 加入 LCD 顯示 → 顯示當前角度和光照強度\n'
    '  4. 雙軸控制 → 水平 + 垂直兩軸同時追蹤',
    'E8F4FD', '003366')

add_box(doc, '完成!',
    '恭喜你完成追日電路！\n'
    '現在太陽爐已經可以自動追蹤太陽，讓發電效率大幅提升！\n'
    '把電路和第 4 章的轉動載臺組合在一起，就是完整的自動追日系統！',
    'D4EDDA', '155724')

# ============================================================
# 附錄
# ============================================================
doc.add_page_break()
add_h1(doc, '附錄：LM393 腳位對照表')

code_block(doc,
    '  LM393 雙比較器 IC 腳位：\n\n'
    '    ┌────────────────────────────────┐\n'
    '    │  ┌───┐                         │\n'
    '    │  │ 1 │  Output A        VCC │8│ ← +5V\n'
    '    │  │ 2 │  Input A-    Output B│7│\n'
    '    │  │ 3 │  Input A+    Input B-│6│\n'
    '    │  │ 4 │  GND         Input B+│5│\n'
    '    │  └───┘                         │\n'
    '    └────────────────────────────────┘\n\n'
    '  注意：\n'
    '  - LM393 是「集極開路」輸出，需要外接上拉電阻到 5V\n'
    '  - 如果用 LM358 運放，則不需要上拉電阻\n'
    '  - 兩者的邏輯一樣：IN+ > IN- → 輸出高；IN+ < IN- → 輸出低'
)

doc.add_paragraph('')
add_h1(doc, '附錄：L293D 腳位對照表')

code_block(doc,
    '  L293D 馬達驅動 IC 腳位：\n\n'
    '    ┌────────────────────────────────┐\n'
    '    │  ┌───┐                         │\n'
    '    │  │ 1 │  Enable 1        VCC │16│ ← +5V (邏輯)\n'
    '    │  │ 2 │  Input 1      Input 4│15│\n'
    '    │  │ 3 │  Output 1     Output 4│14│\n'
    '    │  │ 4 │  GND             GND │13│ ← 接地\n'
    '    │  │ 5 │  GND             GND │12│ ← 接地\n'
    '    │  │ 6 │  Output 2     Output 3│11│\n'
    '    │  │ 7 │  Input 2      Input 3│10│\n'
    '    │  │ 8 │  VS          Enable 2│ 9│ ← +6V~12V (馬達電源)\n'
    '    │  └───┘                         │\n'
    '    └────────────────────────────────┘\n\n'
    '  雙通道用法：\n'
    '  - Channel 1: Pin 1,2,3,6,7 控制馬達 1\n'
    '  - Channel 2: Pin 9,10,11,14,15 控制馬達 2\n'
    '  - 兩個通道可分別驅動水平和垂直馬達'
)

# ============================================================
# 儲存
# ============================================================
output_path = r'G:\我的雲端硬碟\solar-furnace\tools\科展-太陽爐自動追日系統\追日電路設計說明.docx'
doc.save(output_path)
print('Word 檔已儲存至：' + output_path)
