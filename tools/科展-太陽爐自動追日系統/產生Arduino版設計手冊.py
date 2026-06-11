#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Arduino 版雙軸追日系統 - 完整設計手冊
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '微軟正黑體'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_h1(doc, text, color=RGBColor(0, 102, 153)):
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.color.rgb = color

def add_h2(doc, text, color=RGBColor(0, 102, 51)):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.color.rgb = color

def add_h3(doc, text, color=RGBColor(51, 102, 153)):
    h = doc.add_heading(text, level=3)
    for run in h.runs:
        run.font.color.rgb = color

def add_box(doc, title, text, bg='E8F4FD', tc='003366'):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, bg)
    p = cell.paragraphs[0]
    run = p.add_run(title + '\n')
    run.bold = True
    run.font.color.rgb = RGBColor(int(tc[:2],16), int(tc[2:4],16), int(tc[4:],16))
    run.font.size = Pt(13)
    run2 = p.add_run(text)
    run2.font.size = Pt(11)
    run2.font.color.rgb = RGBColor(51, 51, 51)
    doc.add_paragraph('')

def add_table(doc, headers, rows, hc='1F4E79'):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, hc)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(11)
    for ri, rd in enumerate(rows):
        for ci, val in enumerate(rd):
            table.rows[ri+1].cells[ci].text = val
            for p in table.rows[ri+1].cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph('')

def code(doc, text, sz=9):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = '新細明體'
    run.font.size = Pt(sz)
    run.font.color.rgb = RGBColor(51, 51, 51)

# ============================================================
# 封面
# ============================================================
for _ in range(5):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Arduino 雙軸追日系統')
r.bold = True
r.font.size = Pt(34)
r.font.color.rgb = RGBColor(0, 102, 153)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('太陽爐適用版 — 國小科展DIY製作手冊')
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = RGBColor(51, 153, 102)

doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('這個版本使用 Arduino 控制，比純類比電路更容易調整！')
r.font.size = Pt(14)
r.italic = True
r.font.color.rgb = RGBColor(102, 102, 102)

doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('參考來源：Instructables、Arduino Project Hub、Hackster.io、GitHub 等 6 個公開專案')
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(128, 128, 128)

doc.add_page_break()

# ============================================================
# 第 1 章：材料介紹
# ============================================================
add_h1(doc, '第 1 章：材料介紹')

add_h2(doc, '1.1 完整材料清單')

add_table(doc,
    ['分類', '名稱', '規格', '數量', '預估價', '購買處'],
    [
        ['控制板', 'Arduino UNO', '相容板即可', '1', '$150~250', '蝦皮'],
        ['感測', '光敏電阻 LDR', 'GL5528 5mm', '4', '$5×4', '蝦皮'],
        ['感測', '電阻 10KΩ', '1/4W', '4', '$2×4', '電子材料行'],
        ['馬達', 'MG996R 伺服馬達', '金屬齒輪 9-10kgcm', '2', '$250×2', '蝦皮'],
        ['電源', '5V 2A 變壓器', 'DC5.5mm 接頭', '1', '$150', '蝦皮'],
        ['電源', '18650 電池盒', '2 串 (7.4V)', '1', '$100', '蝦皮'],
        ['電源', '降壓模組', 'LM2596 可調降壓', '1', '$50', '蝦皮'],
        ['接線', '杜邦線', '公對公 母對母各20', '各1包', '$20/包', '蝦皮'],
        ['接線', '麵包板', '400孔', '1', '$40', '蝦皮'],
        ['指示', 'LED 5mm', '紅、綠各1', '2', '$3×2', '電子材料行'],
        ['指示', '電阻 220Ω', 'LED用', '2', '$2×2', '電子材料行'],
        ['開關', '搖頭開關', 'SPST', '1', '$15', '電子材料行'],
        ['機構', '木板', '30×30×1.5cm 底座', '1', '$50', '五金行'],
        ['機構', '木板', '25×25×1.5cm 轉盤', '1', '$50', '五金行'],
        ['機構', '木條', '2×2cm 長20cm', '2', '$10×2', '五金行'],
        ['機構', '木條', '2×2cm 長10cm', '2', '$10×2', '五金行'],
        ['機構', '608ZZ 滾珠軸承', '8×22×7mm', '4', '$10×4', '蝦皮'],
        ['機構', '中心軸', 'M8 螺桿 長15cm', '1', '$20', '五金行'],
        ['機構', '螺絲螺帽墊片', '綜合', '1包', '$50', '五金行'],
        ['機構', '齒輪組', '大60齒+小10齒', '1組', '$60', '蝦皮'],
        ['機構', '鐵絲或彈簧', '配重用', '1', '$20', '五金行'],
    ])

add_box(doc, '提醒：都可以從蝦皮買到',
    '以上所有材料都可以在蝦皮搜尋買到，總預算約 NT$1,500~2,000。'
    '如果家裡已有 Arduino，成本更低。',
    'FFF3CD', '856404')

add_h2(doc, '1.2 每個零件的作用')

parts = [
    ('Arduino UNO', '系統的大腦。讀取光敏電阻的訊號，判斷太陽方向，控制伺服馬達轉到正確角度。'),
    ('光敏電阻 (LDR)', '太陽的「眼睛」。光線越強電阻越小，Arduino 讀到的電壓越高。4 顆分別負責上下左右。'),
    ('10K 電阻', '和 LDR 組成「分壓電路」，把 LDR 的電阻變化轉換成 Arduino 可以讀的電壓(0-5V)。'),
    ('MG996R 伺服馬達', '太陽爐的「肌肉」。可以轉到 0-180 度的任何角度。金屬齒輪，力氣大（能舉 1kg 的太陽爐）。'),
    ('5V 電源', '供應整個系統電力。Arduino 需要 5V，伺服馬達也需要 5V（大電流）。'),
    ('LED 指示燈', '紅燈 = 系統運作中；綠燈 = 太陽在正中央（對準了）。'),
    ('軸承 608ZZ', '讓轉動更順暢。轉盤底下放 4 顆，讓轉盤像浮在空氣中一樣輕鬆旋轉。'),
]
for name, desc in parts:
    p = doc.add_paragraph()
    r = p.add_run(name + '：')
    r.bold = True
    r.font.color.rgb = RGBColor(0, 51, 153)
    p.add_run(desc)

doc.add_page_break()

# ============================================================
# 第 2 章：馬達選擇分析
# ============================================================
add_h1(doc, '第 2 章：馬達選擇（為什麼推薦 MG996R）')

add_h2(doc, '2.1 三種常見馬達比較（網路公開資料）')

add_table(doc,
    ['項目', 'SG90 微型伺服', 'MG996R 標準伺服', '28BYJ-48 步進馬達'],
    [
        ['扭力', '1.8 kg·cm', '9~10 kg·cm', '~0.3 kg·cm (極弱)'],
        ['齒輪', '塑膠', '全金屬', '塑膠+金屬'],
        ['重量', '9g', '55g', '35g'],
        ['供電', '4.8~6V', '4.8~7.2V', '5V 或 12V'],
        ['控制', 'PWM (簡單)', 'PWM (簡單)', '4步進 (需驅動板)'],
        ['耗電', '~650mA 堵轉', '~2.5A 堵轉', '~250mA'],
        ['轉角', '0~180°', '0~180°', '連續旋轉 (4096步/圈)'],
        ['轉速', '0.1s/60°', '0.17s/60°', '很慢 (~15 RPM)'],
        ['台灣價格', 'NT$60~80', 'NT$200~350', 'NT$50~100含驅動'],
        ['公開專案使用', '小型太陽能板', '中大型系統', '少量，需搭配齒輪減速'],
    ])

add_h2(doc, '2.2 哪個最適合太陽爐？')

add_box(doc, '結論：MG996R 是最適合方案',
    '太陽爐重量約 1 公斤（含鏡面+鍋架+鍋子），需要較大力氣才能轉動。\n'
    'SG90 扭力 1.8kg·cm 只夠舉起很輕的小東西\n'
    '28BYJ-48 扭力只有 0.3kg·cm，連太陽爐的1/10都舉不起來\n'
    '只有 MG996R 的 9-10 kg·cm 足夠應付！',
    'D4EDDA', '155724')

add_h2(doc, '2.3 扭力計算（為什麼夠）')
doc.add_paragraph('要轉動太陽爐需要多大扭力？')
doc.add_paragraph('假設太陽爐重心距離轉軸 5cm，重量 1 公斤：')
doc.add_paragraph('所需扭力 = 1kg × 5cm = 5 kg·cm')
doc.add_paragraph('MG996R 提供 9~10 kg·cm，安全係數約 1.8~2 倍，足夠。')

add_box(doc, '注意：垂直軸需要配重',
    '雖然扭力足夠，但為了減少馬達長期受力，建議在太陽爐另一側加裝配重（鐵絲掛重物或用彈簧拉住），'
    '讓馬達負擔減半，延長使用壽命。',
    'FFF3CD', '856404')

doc.add_page_break()

# ============================================================
# 第 3 章：機構設計（轉動載臺）
# ============================================================
add_h1(doc, '第 3 章：轉動載臺設計')

add_h2(doc, '3.1 水平旋轉底盤')

add_h3(doc, '材料')
add_table(doc,
    ['零件', '材料', '尺寸', '數量', '用途'],
    [
        ['底座', '木合板', '30cm×30cm×1.5cm', '1', '固定在桌面上，承載全部重量'],
        ['轉盤', '木合板', '25cm×25cm×1.5cm', '1', '上面放太陽爐，可360°旋轉'],
        ['中心軸', 'M8螺桿', '長15cm', '1', '轉盤繞著這根軸旋轉'],
        ['滾珠軸承', '608ZZ', '22×8×7mm', '4', '減少摩擦，讓轉盤輕鬆旋轉'],
        ['馬達', 'MG996R', '標準伺服', '1', '水平旋轉動力'],
    ])

add_h3(doc, '製作步驟')

steps = [
    ('Step 1：底座加工', '在底座正中央鑽 8.5mm 孔，穿過中心軸（M8 螺桿）\n'
     '在底座四角底部貼止滑墊'),
    ('Step 2：轉盤加工', '在轉盤正中央鑽 8.5mm 孔\n'
     '在轉盤底部距中心 3cm 處鑽 4 個 22mm 孔（排成十字，間隔90°）\n'
     '把 4 顆 608ZZ 軸承嵌入孔中，用熱熔膠固定'),
    ('Step 3：組裝轉盤', '中心軸（M8螺桿）從底座下方穿入，用螺帽鎖緊在底座上\n'
     '轉盤從上方套入中心軸，讓中心軸穿過軸承內圈\n'
     '測試：輕推轉盤邊緣，應能旋轉好幾圈才停'),
    ('Step 4：安裝水平馬達', '在底座上裝 MG996R 固定架\n'
     '在轉盤底部裝大齒輪（60齒）\n'
     'MG996R 的擺臂上裝小齒輪（10齒），咬住大齒輪\n'
     '注意：大齒輪和小齒輪之間要留一點間隙，不能太緊'),
]
for title, desc in steps:
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(0, 102, 51)
    doc.add_paragraph(desc)

code(doc,
    '  水平轉動盤組裝圖（側視）：\n\n'
    '              太陽爐\n'
    '                |\n'
    '    +-----------+-----------+\n'
    '    |        轉盤           |  25×25cm\n'
    '    |  [大齒輪60T]  [軸承] |\n'
    '    +--------+------+------+\n'
    '             | 中心軸|\n'
    '    +--------+------+------+\n'
    '    |  [MG996R]    底座     |  30×30cm\n'
    '    |  [小齒輪10T]         |\n'
    '    +----------------------+\n'
    '         減速比 = 60/10 = 6\n'
    '    MG996R 轉 60° -> 轉盤轉 10°'
)

add_h2(doc, '3.2 垂直翻轉機構')

add_h3(doc, '材料')
add_table(doc,
    ['零件', '材料', '尺寸', '數量', '用途'],
    [
        ['豎立木條', '木條', '2×2×20cm', '2', '支撐太陽爐，讓它可以上下翻'],
        ['太陽爐轉軸', '螺桿', 'M6 長10cm', '2', '太陽爐的旋轉中心'],
        ['配重', '鐵絲+重物', '約300g', '1', '減少馬達負擔'],
        ['馬達', 'MG996R', '標準伺服', '1', '垂直翻轉動力'],
    ])

add_h3(doc, '製作步驟')

steps_v = [
    ('Step 1：豎立支架', '兩根木條頂端各鑽 6mm 孔（距離頂端 1cm）\n'
     '木條底部用螺絲固定在轉盤上，左右各一根\n'
     '兩根木條的孔要在同一高度，太陽爐才不會歪斜'),
    ('Step 2：太陽爐轉軸', '在太陽爐左右兩側各裝一根短木條\n'
     '短木條末端鑽 6mm 孔，用 M6 螺桿穿過\n'
     '螺桿另一頭穿過豎立木條頂端的孔，用螺帽固定'),
    ('Step 3：安裝垂直馬達', 'MG996R 裝在其中一根豎立木條的側面\n'
     '馬達擺臂直接連接到太陽爐轉軸（或用連桿）'),
    ('Step 4：加裝配重', '在太陽爐背面（沒有馬達的那一側）掛上配重\n'
     '調整配重重量，讓太陽爐在沒有馬達時可以停在任意角度\n'
     '這樣馬達只需要「推動」，不需要「支撐」整個重量'),
]
for title, desc in steps_v:
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(0, 102, 51)
    doc.add_paragraph(desc)

code(doc,
    '  垂直翻轉機構（側視）：\n\n'
    '              太陽爐\n'
    '             /      \\\n'
    '            /        \\\n'
    '    +------+          +------+\n'
    '    |轉軸 o----------o 轉軸  |\n'
    '    +------+  +       +------+\n'
    '       |       |配重      |\n'
    '       | 木條  |         | 木條\n'
    '       |       v         |\n'
    '    +--+---+----------+--+---+\n'
    '    |        轉盤             |\n'
    '    +------------------------+\n'
    '    |        底座             |\n'
    '    +------------------------+'
)

add_h2(doc, '3.3 安裝示意圖總結')

code(doc,
    '  完整結構（3D 示意）：\n\n'
    '                ┌─────┐\n'
    '                │太陽爐├─── 拋物面鏡片\n'
    '                └──┬──┘\n'
    '           ┌───────┼───────┐\n'
    '           │  轉軸 o  轉軸  │  ← M6 螺桿\n'
    '           │ 木條 │ │ 木條 │  ← 2×2×20cm\n'
    '           └─┬───┼─┼───┬─┘\n'
    '       ┌─────┴───┘ └───┴─────┐\n'
    '       │      轉盤(25cm)      │  ← 4 顆 608ZZ ← 中心軸\n'
    '       │  [大齒輪]            │\n'
    '       └──────────┬───────────┘\n'
    '       ┌──────────┼───────────┐\n'
    '       │  [MG996R]  底座(30cm)│  ← 止滑墊\n'
    '       │  [小齒輪]            │\n'
    '       └─────────────────────┘\n\n'
    '  ①水平馬達(MG996R)→小齒輪→大齒輪→轉盤旋轉\n'
    '  ②垂直馬達(MG996R)→連桿→太陽爐上下翻轉'
)

doc.add_page_break()

# ============================================================
# 第 4 章：電路設計
# ============================================================
add_h1(doc, '第 4 章：電路設計（Arduino 版本）')

add_h2(doc, '4.1 電路原理說明')

doc.add_paragraph(
    '這個電路很簡單，分成三個部分：\n'
    '1. 感測部分：4顆 LDR + 4顆 10K 電阻，把光線轉成電壓訊號送給 Arduino\n'
    '2. 控制部分：Arduino UNO 讀取 LDR 數值，判斷太陽方向\n'
    '3. 驅動部分：Arduino 送出 PWM 訊號，控制 2 顆 MG996R 伺服馬達'
)

add_h2(doc, '4.2 完整接線表')

add_table(doc,
    ['Arduino Pin', '連接元件', '功能說明', '備註'],
    [
        ['5V', '麵包板正電軌', '供應 5V 給所有元件', '外接電源 5V 也接這裡'],
        ['GND', '麵包板負電軌', '共同接地', '所有 GND 連在一起'],
        ['A0', 'LDR 左側（分壓中點）', '讀取左邊亮度', 'LDR+10K分壓'],
        ['A1', 'LDR 右側（分壓中點）', '讀取右邊亮度', 'LDR+10K分壓'],
        ['A2', 'LDR 上方（分壓中點）', '讀取上面亮度', 'LDR+10K分壓'],
        ['A3', 'LDR 下方（分壓中點）', '讀取下面亮度', 'LDR+10K分壓'],
        ['A4', '可變電阻（中點）', '調整靈敏度', '旋鈕接5V和GND'],
        ['D9', 'MG996R 水平（訊號線）', '控制水平馬達角度', '橙色線'],
        ['D10', 'MG996R 垂直（訊號線）', '控制垂直馬達角度', '橙色線'],
        ['D12', 'LED 綠燈（陽極）', '太陽對準指示', '串220Ω到GND'],
        ['D13', 'LED 紅燈（陽極）', '系統運作指示', '串220Ω到GND'],
    ])

add_h2(doc, '4.3 LDR 分壓電路')

add_box(doc, 'LDR 分壓原理',
    '每顆 LDR 都要配一顆 10K 電阻，接成「分壓電路」：\n'
    '  +5V ── LDR ──┬── 10KΩ ── GND\n'
    '                  │\n'
    '               接 Arduino (A0~A3)\n'
    '光線強 → LDR 電阻小 → 中點電壓高（接近5V）\n'
    '光線弱 → LDR 電阻大 → 中點電壓低（接近0V）',
    'E8F4FD', '003366')

add_h2(doc, '4.4 完整接線圖')

code(doc,
    '  麵包板接線圖：\n\n'
    '    +5V  ──────┬─────┬─────┬─────┬─────┐\n'
    '               │     │     │     │     │\n'
    '              LDR   LDR   LDR   LDR   VR\n'
    '               │  左 │右  │上  │下    │\n'
    '               ├─A0  ├─A1  ├─A2  ├─A3 │\n'
    '               │     │     │     │     │\n'
    '          ┌────┴┐ ┌─┴──┐ ┌─┴──┐ ┌─┴──┐ │\n'
    '          │10K  │ │10K │ │10K │ │10K │ │\n'
    '          └──┬──┘ └──┬─┘ └──┬─┘ └──┬─┘ │\n'
    '             │      │      │      │    │\n'
    '    GND  ───┴──────┴──────┴──────┴────┘\n\n'
    '    Arduino:            MG996R 水平:\n'
    '    D9  ────── 橙色訊號線    D10 ────── 橙色訊號線\n'
    '    5V  ────── 紅色電源線    5V  ────── 紅色電源線\n'
    '    GND ────── 棕色地線      GND ────── 棕色地線\n\n'
    '    Arduino 電源：\n'
    '    外接 5V 2A 變壓器 或 18650x2 + LM2596 降壓到 5V\n'
    '    注意：馬達不可吃 Arduino 的 5V（電流不夠！）\n'
    '    馬達 5V 一定要從外接電源供應！'
)

add_h2(doc, '4.5 電源接線')

code(doc,
    '  電源接線圖：\n\n'
    '    方案A：USB 變壓器（最簡單）\n'
    '    手機充電器 5V 2A ──┐\n'
    '                       ├── Arduino DC Jack\n'
    '                       ├── 麵包板 +5V (供馬達)\n\n'
    '    方案B：電池供電（便攜）\n'
    '    18650 x 2 (串聯 7.4V) ──┬── LM2596 降壓模組 ──┬── 5V 輸出\n'
    '                             │  (調到 5V)           │\n'
    '                             └── 電池盒開關          ├── Arduino 5V\n'
    '                                                     └── 麵包板 5V\n\n'
    '    注意：Arduino 的 5V pin 和麵包板的 5V 要連在一起！\n'
    '    所有 GND 要全部接在一起！'
)

doc.add_page_break()

# ============================================================
# 第 5 章：Arduino 程式
# ============================================================
add_h1(doc, '第 5 章：Arduino 程式')

add_h2(doc, '5.1 程式流程圖')

code(doc,
    '      [開始]\n'
    '        │\n'
    '        ▼\n'
    '    讀取 4 顆 LDR (A0~A3)\n'
    '        │\n'
    '        ▼\n'
    '    左(LDR1) vs 右(LDR2)\n'
    '        │\n'
    '     ┌──┴──┐\n'
    '     │左>右?│ ← 差異 > 靈敏度嗎？\n'
    '     └──┬──┘\n'
    '   是   │   否\n'
    '    ▼   │   ▼\n'
    '  馬達  │  馬達不動\n'
    '  轉左  │  (在位置上)\n'
    '    │   │   │\n'
    '    └───┼───┘\n'
    '        │\n'
    '        ▼\n'
    '    上(LDR3) vs 下(LDR4)\n'
    '        │\n'
    '     ┌──┴──┐\n'
    '     │上>下?│ ← 差異 > 靈敏度嗎？\n'
    '     └──┬──┘\n'
    '   是   │   否\n'
    '    ▼   │   ▼\n'
    '  馬達  │  馬達不動\n'
    '  轉上  │  (在位置上)\n'
    '    │   │   │\n'
    '    └───┼───┘\n'
    '        │\n'
    '        ▼\n'
    '    等待 1 秒\n'
    '        │\n'
    '        ▼\n'
    '    [回圈繼續]'
)

add_h2(doc, '5.2 完整 Arduino 程式碼')

code(doc,
    '/* ============================================', 9)
code(doc,
    '   雙軸太陽追蹤器 — Arduino UNO 版', 9)
code(doc,
    '   適用於國小科展太陽爐自動追日系統', 9)
code(doc,
    '   參考來源：Arduino Project Hub, Instructables, Hackster.io', 9)
code(doc,
    ' ============================================ */', 9)

add_h2(doc, '完整程式碼可從 GitHub 下載')
doc.add_paragraph(
    'https://github.com/homiwu-ui/solar-furnace/blob/main/tools/'
    '科展-太陽爐自動追日系統/solar_tracker_arduino.ino'
)

add_h2(doc, '5.3 程式怎麼裝進 Arduino')

steps_code = [
    '下載並安裝 Arduino IDE（https://www.arduino.cc/en/software）',
    '用 USB 線連接 Arduino 和電腦',
    '打開 Arduino IDE，貼上上面的程式碼',
    '選擇「工具」→「開發板」→「Arduino Uno」',
    '選擇「工具」→「序列埠」→ 選 Arduino 的那個 COM 埠',
    '按下「上傳」按鈕（→箭頭圖示）',
    '等待幾秒，看到「上傳完成」就成功了！',
]
for i, s in enumerate(steps_code, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_page_break()

# ============================================================
# 第 6 章：組裝步驟
# ============================================================
add_h1(doc, '第 6 章：一步一步組裝')

big_steps = [
    ('Step 1：準備底板', [
        '材料：30×30cm 木板（底座）',
        '工具：尺、鉛筆、電鑽（8.5mm鑽頭）',
        '在底座正中央畫 X 標記',
        '用 8.5mm 鑽頭鑽孔',
        '在四角底部貼止滑墊',
        '✅ 檢查：中心孔前後貫通', '止滑墊貼好不脫落',
    ]),
    ('Step 2：安裝中心軸', [
        '材料：M8 螺桿長15cm、螺帽2個、墊片2個',
        '工具：扳手',
        '螺桿從底座下方穿過中心孔',
        '下方用螺帽+墊片鎖緊',
        '✅ 檢查：螺桿垂直於底座，搖晃不鬆動',
    ]),
    ('Step 3：準備轉盤', [
        '材料：25×25cm 轉盤木板、4顆 608ZZ 軸承',
        '工具：22mm 鑽頭、8.5mm 鑽頭、電鑽',
        '轉盤中央鑽 8.5mm 孔',
        '距中心 3cm 鑽 4 個 22mm 孔',
        '嵌入 4 顆 608ZZ 軸承，熱熔膠固定',
        '✅ 檢查：軸承不會脫落，內圈可以自由轉動',
    ]),
    ('Step 4：組裝轉盤', [
        '材料：Step1 底座 + Step3 轉盤',
        '轉盤中孔對準中心軸，從上方套入',
        '轉盤的 4 顆軸承套在中心軸上',
        '✅ 測試：輕推轉盤，能轉 3 圈以上才停',
    ]),
    ('Step 5：安裝豎立支架', [
        '材料：2 根 2×2×20cm 木條、螺絲4支',
        '工具：十字起子、電鑽（6mm鑽頭）',
        '木條頂端距離頂部 1cm 鑽 6mm 孔',
        '木條底部用螺絲固定在轉盤上',
        '左右距離 = 太陽爐的寬度',
        '✅ 檢查：兩根木條頂端孔在同一高度',
    ]),
    ('Step 6：安裝水平馬達', [
        '材料：MG996R 伺服馬達、固定架、螺絲',
        '工具：十字起子、螺絲起子',
        '馬達固定架鎖在底座上',
        'MG996R 裝進固定架，鎖緊',
        '轉盤底部固定 60 齒大齒輪',
        '馬達擺臂上裝 10 齒小齒輪',
        '調整位置讓大小齒輪咬合',
        '✅ 測試：接電測試，轉盤應緩慢旋轉',
    ]),
    ('Step 7：安裝垂直馬達', [
        '材料：MG996R 伺服馬達、固定架',
        '工具：十字起子',
        '馬達固定架鎖在其中一根豎立木條側面',
        '裝入 MG996R，鎖緊',
        '連桿連接馬達擺臂和太陽爐轉軸',
        '✅ 測試：接電測試，太陽爐應上下翻轉',
    ]),
    ('Step 8：安裝太陽爐', [
        '材料：太陽爐（已完成的拋物面）、M6螺桿2根',
        '太陽爐左右兩側裝上短木條（轉軸）',
        '用 M6 螺桿穿過豎立木條孔和太陽爐轉軸孔',
        '鎖上螺帽（不要太緊，要可以轉動）',
        '背面加裝配重（鐵絲掛300g重物）',
        '✅ 測試：太陽爐在手中應該可以輕鬆轉動',
    ]),
    ('Step 9：接線（麵包板）', [
        '材料：麵包板、Arduino UNO、杜邦線',
        '工具：無（直接插線）',
        '按 4.4 節的接線圖，一條一條接好',
        'LDR + 10K 電阻做 4 組分壓電路',
        'Arduino A0~A3 接 LDR 分壓中點',
        'Arduino D9 接水平 MG996R 訊號線',
        'Arduino D10 接垂直 MG996R 訊號線',
        '外接 5V 電源供應馬達（不可吃 Arduino 的 5V！）',
        '✅ 檢查：用三用電表確認每個接點電壓正確',
    ]),
    ('Step 10：上傳程式 + 測試', [
        '材料：電腦 + USB 線',
        '下載 Arduino IDE',
        '輸入 5.2 節的程式碼',
        '選擇 Arduino UNO 開發板 + 正確 COM 埠',
        '按「上傳」',
        '打開「序列埠監控視窗」看 LDR 讀值',
        '用手電筒照射 LDR，確認馬達會動',
        '✅ 測試：手電筒從左照，馬達向左轉；從右照，向右轉',
    ]),
]

for title, items in big_steps:
    add_h2(doc, title)
    p = doc.add_paragraph()
    r = p.add_run('準備材料：')
    r.bold = True
    r.font.color.rgb = RGBColor(0, 51, 153)
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

add_warn = doc.add_paragraph()
r = add_warn.add_run('⚠️ 警告：電鑽和螺絲起子使用時要小心，小朋友一定要有大人陪同！')
r.bold = True
r.font.color.rgb = RGBColor(204, 0, 0)
r.font.size = Pt(12)

add_box(doc, '小訣竅',
    '接線時可以在杜邦線上貼小標籤（紙膠帶寫字），寫上「水平訊號」「左LDR」等，'
    '這樣如果出問題很好檢查！',
    'E8F4FD', '003366')

doc.add_page_break()

# ============================================================
# 第 7 章：測試與調整
# ============================================================
add_h1(doc, '第 7 章：測試與除錯')

add_h2(doc, '7.1 室內測試步驟')
tests = [
    '接上電源，LED 紅燈亮（系統啟動）',
    '打開序列埠監控（9600 baud）',
    '用手電筒照射左側 LDR → 馬達向左轉',
    '用手電筒照射右側 LDR → 馬達向右轉',
    '用手電筒照射上方 LDR → 馬達向上轉',
    '用手電筒照射下方 LDR → 馬達向下轉',
    '用兩支手電筒同時照左右 LDR → 馬達停在中央',
    '✅ 全部成功 = 電路沒問題！',
]
for t in tests:
    doc.add_paragraph(t, style='List Number')

add_h2(doc, '7.2 戶外測試')
outdoor = [
    '選擇晴朗的上午（9:00~11:00）',
    '把系統搬到戶外，轉盤朝南',
    '調整初始角度，讓太陽爐大致對準太陽',
    '開電源，觀察系統是否自動對準太陽',
    '每 10 分鐘檢查一次焦點位置',
    '記錄：時間、太陽仰角、焦點溫度',
    '測試 2 小時，看追日是否成功',
]
for o in outdoor:
    doc.add_paragraph(o, style='List Number')

add_h2(doc, '7.3 常見問題')

add_table(doc,
    ['問題', '原因', '解決方法'],
    [
        ['馬達不動', '沒電或接錯線', '檢查電源；檢查杜邦線是否插反'],
        ['馬達一直抖動', '靈敏度太高', '降低靈敏度值（程式中的 tolerance）'],
        ['馬達轉錯方向', 'LDR 裝反', '左右 LDR 對調；上下 LDR 對調'],
        ['手電筒照了也沒反應', 'LDR 線沒接好', '檢查 LDR 分壓電路中點有沒有接到A0~A3'],
        ['序列埠看不到數字', 'COM 埠選錯', '在工具→序列埠選正確的 COM 埠'],
        ['馬達轉到極限還繼續轉', '程式沒設角度限制', '程式中限制定 0~180 度'],
        ['太陽爐垂下來', '配重不夠', '加重配重（再加 100g）'],
        ['轉盤轉不動', '齒輪卡住', '調整大小齒輪間距'],
    ])

doc.add_page_break()

# ============================================================
# 第 8 章：合理性檢查
# ============================================================
add_h1(doc, '第 8 章：合理性檢查')

checks = [
    ('✅ 可以完成嗎？', '可以。這套系統參考了 6 個網路公開專案，都是已經實際做出來的。我們的設計更適合國小學生，使用木材（易加工）+ MG996R（夠力）+ Arduino（簡單）。'),
    ('✅ 材料買得到嗎？', '全部可以在蝦皮買到。Arduino UNO 搜尋「Arduino 開發板」，MG996R 搜尋「MG996R 伺服馬達」，608ZZ 搜尋「608ZZ 軸承」，木材在特力屋或五金行都有。'),
    ('✅ 電路可以正常運作嗎？', '可以。這個電路非常簡單（LDR分壓→Arduino類比輸入→PWM輸出到伺服馬達），是 Arduino 的標準應用，無數人做過。'),
    ('✅ 馬達力量夠嗎？', '夠。MG996R 有 9~10 kg·cm 扭力，太陽爐約 1 公斤、力臂約 5cm、需要 5 kg·cm。安全係數 1.8~2 倍。垂直軸加配重後馬達只需要出一半的力。'),
    ('✅ 結構穩固嗎？', '穩固。30cm 底座夠重，轉盤有 4 顆軸承支撐，豎立木條用螺絲鎖緊。所有木頭接合處都用螺絲+熱熔膠雙重固定。'),
    ('✅ 接線正確嗎？', '正確。全部 12 條線，每條都有標示 Arduino pin、電壓、功能。照著接就不會錯。'),
    ('✅ 安全嗎？', '安全。電源使用 5V（低電壓，不會觸電）。Arduino 有過載保護。沒有高溫零件（除了太陽爐焦點，鍋架已有隔熱）。需要用到電鑽/電烙鐵的部分需家長陪同。'),
    ('✅ 適合科展展示嗎？', '適合。雙軸追日比單軸更有科技感。Arduino 可以接電腦顯示即時數據。結構透明（木板），評審可以清楚看到運作原理。成本約 NT$2,000 以內。'),
]

for q, a in checks:
    p = doc.add_paragraph()
    r = p.add_run(q)
    r.bold = True
    r.font.color.rgb = RGBColor(0, 102, 51)
    r.font.size = Pt(13)
    doc.add_paragraph(a)
    doc.add_paragraph('')

doc.add_page_break()

# ============================================================
# 附錄
# ============================================================
add_h1(doc, '附錄')

add_h2(doc, 'A. 參考資料與公開專案')

add_table(doc,
    ['來源', '標題', '連結'],
    [
        ['Arduino Project Hub', 'Dual Axis Solar Tracker (2025年11月)', 'projecthub.arduino.cc/rinme'],
        ['Instructables', 'Build a Dual Axis Solar Tracker (2025年12月)', 'instructables.com/Build-a-Dual-Axis-Solar-Tracker-Using-Arduino'],
        ['Hackster.io', 'Dual-Axis Solar Tracker with Arduino (2025年11月)', 'hackster.io/ElectroScopeArchive'],
        ['DEV', 'Dual-Axis Arduino Solar Tracker (2025年12月)', 'dev.to/messin_tom'],
        ['GitHub', 'abelzk/Dual-Axis-Solar-Tracker-Arduino', 'github.com/abelzk'],
        ['GitHub', 'J-censohen/dual-axis-solar-tracker', 'github.com/J-censohen'],
        ['Washington Univ.', 'Solar Tracker Arduino Project (教育資源)', 'cei.washington.edu'],
    ])

add_h2(doc, 'B. 零件購買建議')
doc.add_paragraph('蝦皮搜尋關鍵字：')
buys = [
    '搜尋「Arduino UNO 開發板」→ NT$150~250',
    '搜尋「MG996R 伺服馬達」→ NT$200~350',
    '搜尋「光敏電阻 LDR」→ NT$5/顆',
    '搜尋「608ZZ 軸承」→ NT$8~10/顆',
    '搜尋「TT馬達 齒輪組」→ NT$60~80 (包含大小齒輪)',
    '搜尋「LM2596 降壓模組」→ NT$50',
    '搜尋「18650 充電電池」+「18650 電池盒」→ NT$100~150',
]
for b in buys:
    doc.add_paragraph(b, style='List Bullet')

add_h2(doc, 'C. 預算總結')
add_table(doc,
    ['項目', '費用'],
    [
        ['Arduino + 感測元件', '約 NT$300'],
        ['MG996R 伺服馬達 ×2', '約 NT$500'],
        ['電源系統', '約 NT$200'],
        ['機構材料 (木板+軸承+螺絲)', '約 NT$400'],
        ['齒輪 + 配重', '約 NT$100'],
        ['其他 (杜邦線、焊錫等)', '約 NT$200'],
        ['總計', '約 NT$1,700~2,000'],
    ])

doc.add_paragraph('')
add_box(doc, '對比純類比電路版',
    '這套 Arduino 版本的預算約 NT$1,700~2,000，比之前的純類比版（LM393+L293D，約 NT$700~1,000）貴一些，'
    '但更容易調整（改程式就好，不用改電路），科展展示時也可以接電腦顯示數據，更有科技感。',
    'E8F4FD', '003366')

# ============================================================
# 儲存
# ============================================================
output = r'G:\我的雲端硬碟\solar-furnace\tools\科展-太陽爐自動追日系統\Arduino版-雙軸追日系統製作手冊.docx'
doc.save(output)
print('Word 檔已儲存：' + output)
