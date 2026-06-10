#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
第 4 章：製作轉動載臺 — 細部說明 Word 檔
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '微軟正黑體'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_h1(doc, text, color=RGBColor(0, 102, 153)):
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.color.rgb = color

def add_h2(doc, text, color=RGBColor(0, 102, 51)):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.color.rgb = color

def add_h3(doc, text, color=RGBColor(51, 102, 153)):
    h = doc.add_heading(text, level=3)
    for run in h.runs:
        run.font.color.rgb = color

def add_box(doc, title, text, bg='E8F4FD', tc='003366'):
    """加入提示/警告框"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, bg)
    p = cell.paragraphs[0]
    run = p.add_run(title + '\n')
    run.bold = True
    run.font.color.rgb = RGBColor(int(tc[:2],16), int(tc[2:4],16), int(tc[4:],16))
    run.font.size = Pt(13)
    run2 = p.add_run(text)
    run2.font.size = Pt(11)
    run2.font.color.rgb = RGBColor(51, 51, 51)
    doc.add_paragraph('')

def add_parts_table(doc, items):
    """零件清單表格"""
    table = doc.add_table(rows=1+len(items), cols=4)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['物品', '規格', '數量', '備註']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, '1F4E79')
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(11)
    for ri, rd in enumerate(items):
        for ci, val in enumerate(rd):
            table.rows[ri+1].cells[ci].text = val
            for p in table.rows[ri+1].cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph('')

def code_block(doc, text):
    """程式碼/圖示區塊"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = '新細明體'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(51, 51, 51)
    return p

# ============================================================
# 封面
# ============================================================
for _ in range(6):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('第 4 章')
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0, 102, 153)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('製作轉動載臺')
r.bold = True
r.font.size = Pt(36)
r.font.color.rgb = RGBColor(0, 102, 153)

doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('讓太陽爐可以「東西轉」和「上下翻」')
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(102, 102, 102)
r.italic = True

doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('本章共 4 節，請依序製作')
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(102, 102, 102)

doc.add_page_break()

# ============================================================
# 章節導覽
# ============================================================
add_h1(doc, '本章目錄')
sections = [
    ('4.1', '水平轉動盤', '讓太陽爐可以左右旋轉（東西向跟隨太陽）'),
    ('4.2', '垂直轉動機構', '讓太陽爐可以上下翻轉（調整仰角）'),
    ('4.3', '安裝馬達與齒輪', '用減速馬達自動控制兩個方向的轉動'),
    ('4.4', '組裝與測試', '把所有零件組合起來並測試轉動'),
]
for num, title, desc in sections:
    p = doc.add_paragraph()
    r = p.add_run(f'{num}  {title}')
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0, 102, 153)
    doc.add_paragraph(f'    {desc}')

doc.add_page_break()

# ============================================================
# 4.1 水平轉動盤
# ============================================================
add_h1(doc, '4.1 水平轉動盤')

add_h2(doc, '設計原理')
doc.add_paragraph(
    '水平轉動盤的任務是讓太陽爐可以「繞著中心軸左右旋轉」，'
    '這樣才能從東方追到西方。'
    '我們需要一個「轉盤」和一個「底座」，中間用軸承連接，'
    '讓轉盤可以輕鬆旋轉。'
)

add_box(doc, '核心概念',
    '想像一張旋轉餐桌：桌面（轉盤）可以繞著中間的柱子（中心軸）輕鬆旋轉。'
    '我們的水平轉動盤就是一樣的道理！',
    'E8F4FD', '003366')

add_h2(doc, '零件清單')
add_parts_table(doc, [
    ['木合板（轉盤）', '1.5cm厚 25x25cm', '1片', '上下兩面都要磨平'],
    ['木合板（底座）', '1.5cm厚 30x30cm', '1片', '底部可加止滑墊'],
    ['滾珠軸承 608ZZ', '8x22x7mm（內徑8mm）', '4顆', '放在轉盤底部中央附近'],
    ['中心軸（木棒或螺桿）', '直徑8mm 長15cm', '1根', '要能穿過軸承內圈'],
    ['木牙螺絲', '3cm', '8支', '固定軸承座'],
    ['熱熔膠', '膠條', '1支', '輔助固定'],
    ['止滑墊', '3x3cm', '4片', '貼在底座四角底部'],
])

add_h2(doc, '詳細製作步驟')

add_h3(doc, 'Step 1：製作底座')
steps_1 = [
    '取 30x30cm 木合板作為底座',
    '在底座四個角落各貼一片止滑墊，防止底座滑動',
    '在底座正中央標記一個點，用 8.5mm 鑽頭鑽一個洞（中心軸孔）',
    '這個洞要能讓中心軸穿過，但不能太鬆',
]
for s in steps_1:
    doc.add_paragraph(s, style='List Number')

code_block(doc,
    '  底座示意圖（上視圖）：\n\n'
    '    +-------------------------------+\n'
    '    | [止滑墊]             [止滑墊] |  30x30cm\n'
    '    |                               |\n'
    '    |                               |\n'
    '    |              o                |  <-- 中心軸孔（8.5mm）\n'
    '    |                               |\n'
    '    |                               |\n'
    '    | [止滑墊]             [止滑墊] |\n'
    '    +-------------------------------+'
)

add_h3(doc, 'Step 2：製作轉盤')
steps_2 = [
    '取 25x25cm 木合板作為轉盤',
    '在轉盤正中央標記一個點，用 8.5mm 鑽頭鑽一個洞（中心軸孔）',
    '在轉盤底部，距離中心 3~5cm 處，標記 4 個位置（間隔 90°）',
    '每個位置用 22mm 鑽頭鑽一個大洞，用來放 608ZZ 軸承',
    '把 4 顆 608ZZ 軸承嵌入這 4 個洞中，用熱熔膠固定',
]
for s in steps_2:
    doc.add_paragraph(s, style='List Number')

code_block(doc,
    '  轉盤底部示意圖（仰視圖）：\n\n'
    '    +-------------------------------+\n'
    '    |                               |  25x25cm\n'
    '    |          +-----+              |\n'
    '    |          |軸承2|              |\n'
    '    |          +--+--+              |\n'
    '    |             |                 |\n'
    '    |  +----------o----------+      |  <-- o = 中心軸孔\n'
    '    |  |軸承1    中心    軸承3|      |\n'
    '    |  +----------+----------+      |\n'
    '    |             |                 |\n'
    '    |          +--+--+              |\n'
    '    |          |軸承4|              |\n'
    '    |          +-----+              |\n'
    '    |                               |\n'
    '    +-------------------------------+\n\n'
    '  4顆軸承排成十字形，距離中心約 3~5cm\n'
    '  軸承內圈（8mm）朝上，用來套住中心軸'
)

add_h3(doc, 'Step 3：安裝中心軸')
steps_3 = [
    '把中心軸（木棒或螺桿）從底座下方穿過中心孔',
    '中心軸要從底座上方突出至少 12cm',
    '在底座下方用螺帽或熱熔膠固定中心軸，讓它不會上下滑動',
    '確保中心軸是垂直的（用目視或水平尺檢查）',
]
for s in steps_3:
    doc.add_paragraph(s, style='List Number')

code_block(doc,
    '  側面圖：\n\n'
    '                    中心軸\n'
    '                      |\n'
    '                      |  <-- 突出 12cm\n'
    '    +-----------------+-----------------+\n'
    '    |               底座               |  <-- 30x30cm\n'
    '    +-----------------+-----------------+\n'
    '                      |\n'
    '                  螺帽固定\n\n'
    '  中心軸從底座下方穿出，用螺帽鎖緊'
)

add_h3(doc, 'Step 4：組合轉盤與底座')
steps_4 = [
    '把轉盤的中心孔對準中心軸，從上方套下去',
    '轉盤的 4 顆軸承內圈會套住中心軸',
    '用手輕輕轉動轉盤，確認它能繞著中心軸順暢旋轉',
    '如果太緊：用砂紙稍微磨一下軸承內圈或中心軸',
    '如果太鬆：在中心軸上纏一圈膠帶增加摩擦力',
]
for s in steps_4:
    doc.add_paragraph(s, style='List Number')

code_block(doc,
    '  組合側面圖：\n\n'
    '    轉盤（25x25cm）\n'
    '    +===========================+\n'
    '    |  o   o   o   o   o   o   |  <-- 4顆軸承（十字排列）\n'
    '    +------------+--------------+\n'
    '                 |\n'
    '    +------------+--------------+\n'
    '    |           底座            |  <-- 30x30cm\n'
    '    +------------+--------------+\n'
    '                 |\n'
    '             止滑墊\n\n'
    '  轉盤透過軸承套在中心軸上，可以輕鬆旋轉'
)

add_box(doc, '小測試',
    '把轉盤放在桌上，用手輕推轉盤邊緣。它應該要能轉好幾圈才停下來。'
    '如果推不動或轉一下就停，檢查軸承有沒有裝好。',
    'FFF3CD', '856404')

doc.add_page_break()

# ============================================================
# 4.2 垂直轉動機構
# ============================================================
add_h1(doc, '4.2 垂直轉動機構')

add_h2(doc, '設計原理')
doc.add_paragraph(
    '垂直轉動機構的任務是讓太陽爐可以「上下翻轉」，'
    '調整仰角來配合太陽在天空中的高度。'
    '太陽早上和傍晚比較低，中午比較高，所以我們需要讓太陽爐可以上下調整角度。'
)

add_box(doc, '核心概念',
    '想像你在看天上的飛機：你會把頭抬高或低下，讓視線對準飛機。'
    '垂直轉動機構就像你的脖子，讓太陽爐「抬頭」或「低頭」。',
    'E8F4FD', '003366')

add_h2(doc, '零件清單')
add_parts_table(doc, [
    ['豎立木條', '2x2cm 20cm長', '2根', '作為垂直支架'],
    ['水平木條', '2x2cm 25cm長', '2根', '連接兩根豎立木條的頂端'],
    ['短木條', '2x2cm 10cm長', '2根', '裝在太陽爐兩側作為轉軸'],
    ['螺桿', '直徑6mm 長15cm', '2根', '作為垂直轉軸'],
    ['螺帽', 'M6', '4顆', '固定轉軸'],
    ['軸承 608ZZ', '8x22x7mm', '2顆', '減少轉動摩擦'],
    ['木牙螺絲', '3cm', '8支', '固定支架'],
])

add_h2(doc, '詳細製作步驟')

add_h3(doc, 'Step 1：製作垂直支架')
steps_v1 = [
    '取兩根 20cm 長的豎立木條',
    '在每根木條的頂端（距離頂端 1cm 處）鑽一個 6mm 的洞',
    '這個洞就是垂直轉軸的穿過點',
    '兩根豎立木條的洞要在同一個高度，這樣太陽爐才不會歪斜',
]
for s in steps_v1:
    doc.add_paragraph(s, style='List Number')

code_block(doc,
    '  豎立木條示意圖：\n\n'
    '    +--o--+     +--o--+      <-- o = 轉軸洞（距離頂端1cm）\n'
    '    |     |     |     |\n'
    '    |     |     |     |      20cm\n'
    '    |     |     |     |\n'
    '    |     |     |     |\n'
    '    +-----+     +-----+\n'
    '    豎立木條A   豎立木條B\n\n'
    '    兩根木條的洞要在同一高度！'
)

add_h3(doc, 'Step 2：製作太陽爐轉軸')
steps_v2 = [
    '在太陽爐左右兩側各裝一根 10cm 短木條',
    '短木條要從太陽爐側面伸出至少 8cm',
    '在每根短木條的末端鑽一個 6mm 的洞',
    '這個洞就是垂直轉軸，會穿過豎立木條頂端的洞',
]
for s in steps_v2:
    doc.add_paragraph(s, style='List Number')

code_block(doc,
    '  太陽爐轉軸示意圖（正面）：\n\n'
    '              太陽爐\n'
    '            /        \\\n'
    '           /          \\\n'
    '    +-----+            +-----+\n'
    '    |短木條|            |短木條|  <-- 各伸出 8cm\n'
    '    +--o--+            +--o--+  <-- o = 轉軸洞\n'
    '       |                  |\n'
    '       |                  |\n'
    '       v                  v\n'
    '    穿過豎立木條的頂端洞'
)

add_h3(doc, 'Step 3：組裝垂直支架')
steps_v3 = [
    '把兩根豎立木條固定在轉盤上（用螺絲從下方鎖入）',
    '兩根木條之間的距離要等於太陽爐的寬度',
    '在兩根豎立木條的頂端各裝一顆軸承（608ZZ），減少轉動摩擦',
    '把太陽爐的轉軸穿過豎立木條頂端的洞和軸承',
    '兩端用螺帽固定，但不要鎖太緊，要讓太陽爐可以自由上下翻轉',
]
for s in steps_v3:
    doc.add_paragraph(s, style='List Number')

code_block(doc,
    '  組裝側面圖：\n\n'
    '                    太陽爐\n'
    '                   /      \\\n'
    '                  /        \\\n'
    '        +--------+          +--------+\n'
    '        | 轉軸 o------------o 轉軸   |\n'
    '        +--------+          +--------+\n'
    '            |                    |\n'
    '            | 豎立木條           | 豎立木條\n'
    '            | (20cm)            | (20cm)\n'
    '            |                    |\n'
    '        +---+--------------------+---+\n'
    '        |          轉盤              |  <-- 水平轉動盤\n'
    '        +----------------------------+\n'
    '        |          底座              |\n'
    '        +----------------------------+\n\n'
    '  太陽爐透過轉軸架在兩根豎立木條之間\n'
    '  可以繞著轉軸上下翻轉'
)

add_h3(doc, 'Step 4：測試垂直轉動')
steps_v4 = [
    '用手輕推太陽爐的邊緣',
    '太陽爐應該可以輕鬆地向上或向下翻轉',
    '測試範圍：從水平（0°）到接近垂直（80°）',
    '如果轉不動：檢查轉軸有沒有卡住，螺帽有沒有鎖太緊',
    '如果太鬆（太陽爐會自己垂下來）：稍微鎖緊螺帽',
]
for s in steps_v4:
    doc.add_paragraph(s, style='List Number')

add_box(doc, '角度小知識',
    '太陽在天空中的仰角：\n'
    '  - 早上/傍晚：約 10~30°\n'
    '  - 上午/下午：約 30~60°\n'
    '  - 中午：約 60~80°（台灣夏季）\n'
    '所以我們的垂直轉動至少要能覆蓋 0°~80° 的範圍。',
    'E8F4FD', '003366')

doc.add_page_break()

# ============================================================
# 4.3 安裝馬達與齒輪
# ============================================================
add_h1(doc, '4.3 安裝馬達與齒輪')

add_h2(doc, '為什麼要用減速馬達？')
doc.add_paragraph(
    '太陽爐重約 1 公斤，如果用一般小馬達（轉速數千RPM），'
    '轉太快會讓太陽爐晃動、不穩定，而且力氣不夠大推不動。'
    '所以我们需要用「減速馬達」——轉速慢但力氣大。'
)

add_box(doc, '重要提醒',
    '一定要用「減速馬達」！\n'
    '  - TT 減速馬達 1:48（推薦）：6V供電，約4RPM，力氣大\n'
    '  - N20 減速馬達 30RPM（升級選擇）：更安靜更精準\n'
    '  - 一般小馬達（不建議）：轉太快，力氣不夠',
    'FFF3CD', '856404')

add_h2(doc, '零件清單')
add_parts_table(doc, [
    ['TT 減速馬達', '6V 1:48', '2顆', '1顆水平、1顆垂直'],
    ['小齒輪（馬達用）', '適配TT馬達軸', '2個', '通常馬達會附'],
    ['大齒輪（轉盤用）', '直徑6~8cm', '1個', '固定在轉盤底部'],
    ['大齒輪（垂直用）', '直徑6~8cm', '1個', '固定在垂直軸上'],
    ['馬達固定架', 'TT馬達專用', '2個', '固定馬達用'],
    ['木牙螺絲', '2cm', '8支', '固定馬達架'],
    ['杜邦線', '公對母', '4條', '連接馬達到電路'],
])

add_h2(doc, '水平馬達安裝')

add_h3(doc, '原理：小齒輪帶動大齒輪')
doc.add_paragraph(
    '馬達軸上裝一個小齒輪（約10齒），轉盤底部裝一個大齒輪（約60齒）。'
    '小齒輪轉很多圈，大齒輪才轉一圈，這樣就能把馬達的高速旋轉'
    '變成慢速但力氣大的旋轉。'
)

code_block(doc,
    '  齒輪減速原理：\n\n'
    '    馬達軸（小齒輪 10齒）     轉盤（大齒輪 60齒）\n'
    '         +---+                     +--------+\n'
    '         | * | ----------------->  |   *    |\n'
    '         +---+   減速比 6:1        +--------+\n'
    '          轉6圈                    才轉1圈\n\n'
    '    減速比 = 大齒輪齒數 / 小齒輪齒數 = 60 / 10 = 6\n'
    '    馬達轉速 240RPM -> 轉盤轉速 240/6 = 40RPM\n'
    '    再乘上馬達內部減速 1:48 -> 最終 240/48/6 ≈ 0.8RPM\n'
    '    也就是約 75 秒才轉一圈，非常穩定！'
)

add_h3(doc, '安裝步驟')
steps_h = [
    '在轉盤底部中央偏邊的位置，固定大齒輪（用熱熔膠或螺絲）',
    '大齒輪要和轉盤一起轉動',
    '在底座上安裝馬達固定架，位置要讓小齒輪能咬住大齒輪',
    '把 TT 馬達裝進固定架，確保小齒輪和大齒輪完全嚙合',
    '用螺絲把馬達固定架鎖在底座上',
    '測試：接上電池，馬達應該能帶動整個轉盤緩慢旋轉',
]
for s in steps_h:
    doc.add_paragraph(s, style='List Number')

code_block(doc,
    '  水平馬達安裝側面圖：\n\n'
    '            轉盤\n'
    '    +===================+\n'
    '    |                   |\n'
    '    |      [大齒輪]     |  <-- 固定在轉盤底部\n'
    '    |          *        |\n'
    '    +----------+--------+\n'
    '               |\n'
    '    +----------+--------+\n'
    '    |       [小齒輪]    |  <-- 固定在馬達軸上\n'
    '    |          *        |\n'
    '    |       [馬達]      |  <-- TT 減速馬達\n'
    '    |                   |\n'
    '    +-------------------+\n'
    '           底座\n\n'
    '  小齒輪咬住大齒輪，馬達轉動時帶動轉盤旋轉'
)

add_h2(doc, '垂直馬達安裝')

add_h3(doc, '安裝步驟')
steps_v = [
    '在其中一根豎立木條的側面安裝馬達固定架',
    '馬達軸要對準垂直轉軸的方向',
    '在垂直轉軸上固定大齒輪（用熱熔膠或鍊條連接）',
    '把 TT 馬達裝進固定架，確保小齒輪能驅動大齒輪',
    '測試：接上電池，馬達應該能帶動太陽爐上下翻轉',
]
for s in steps_v:
    doc.add_paragraph(s, style='List Number')

code_block(doc,
    '  垂直馬達安裝側面圖：\n\n'
    '                    太陽爐\n'
    '                   /      \\\n'
    '                  /        \\\n'
    '        +--------+          +--------+\n'
    '        | 轉軸 o-----*------o 轉軸   |  <-- * = 大齒輪\n'
    '        +--------+          +--------+\n'
    '            |                    |\n'
    '            | 豎立木條           | 豎立木條\n'
    '            |                    |\n'
    '            |  [馬達]            |\n'
    '            |    *               |  <-- 小齒輪\n'
    '            |  [固定架]          |\n'
    '        +---+--------------------+---+\n'
    '        |          轉盤              |\n'
    '        +----------------------------+\n\n'
    '  垂直馬達裝在豎立木條側面，驅動太陽爐上下翻轉'
)

doc.add_page_break()

# ============================================================
# 4.4 組裝與測試
# ============================================================
add_h1(doc, '4.4 組裝與測試')

add_h2(doc, '完整組裝順序')
assembly = [
    ('Step 1', '底座', '貼止滑墊 -> 鑽中心軸孔 -> 裝中心軸'),
    ('Step 2', '轉盤', '鑽中心孔 -> 鑽軸承孔 -> 裝4顆軸承 -> 套入中心軸'),
    ('Step 3', '垂直支架', '鑽轉軸洞 -> 固定豎立木條到轉盤'),
    ('Step 4', '太陽爐轉軸', '裝短木條 -> 穿過豎立木條 -> 螺帽固定'),
    ('Step 5', '水平馬達', '裝大齒輪到轉盤 -> 裝馬達架 -> 裝馬達'),
    ('Step 6', '垂直馬達', '裝大齒輪到轉軸 -> 裝馬達架 -> 裝馬達'),
    ('Step 7', '接線', '馬達接杜邦線 -> 準備連接電路板'),
]
table = doc.add_table(rows=1+len(assembly), cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['步驟', '部位', '動作']):
    cell = table.rows[0].cells[i]
    cell.text = h
    set_cell_shading(cell, '1F4E79')
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
for ri, (step, part, action) in enumerate(assembly):
    table.rows[ri+1].cells[0].text = step
    table.rows[ri+1].cells[1].text = part
    table.rows[ri+1].cells[2].text = action
doc.add_paragraph('')

add_h2(doc, '測試項目')
tests = [
    ('水平旋轉', '用手推轉盤，應能360°順暢旋轉；接上馬達後應能緩慢穩定旋轉'),
    ('垂直翻轉', '用手推太陽爐，應能0°~80°上下翻轉；接上馬達後應能緩慢調整仰角'),
    ('負載測試', '在太陽爐上放一個1公斤的重物（如水杯），測試馬達能否正常推動'),
    ('追日模擬', '用手電筒從不同角度照射LDR，觀察馬達是否會自動轉動'),
]
for name, desc in tests:
    p = doc.add_paragraph()
    r = p.add_run(name + '：')
    r.bold = True
    r.font.color.rgb = RGBColor(0, 102, 51)
    p.add_run(desc)

add_h2(doc, '常見問題')
problems = [
    ('轉盤轉不動', '軸承沒裝好、中心軸太粗、熱熔膠卡住軸承 -> 檢查並重新安裝'),
    ('太陽爐會自己垂下來', '垂直軸螺帽太鬆 -> 稍微鎖緊螺帽，增加摩擦力'),
    ('馬達轉但齒輪不動', '小齒輪和大齒輪沒咬合好 -> 調整馬達位置'),
    ('馬達轉反了', '正負極接反 -> 把杜邦線對調'),
]
for prob, sol in problems:
    p = doc.add_paragraph()
    r = p.add_run('Q：' + prob)
    r.bold = True
    r.font.color.rgb = RGBColor(0, 51, 153)
    doc.add_paragraph('A：' + sol)
    doc.add_paragraph('')

add_box(doc, '完成!',
    '恭喜你完成轉動載臺！現在太陽爐已經可以東西轉動和上下翻轉了。\n'
    '下一步是第 5 章：製作追日電路，讓系統自動追蹤太陽！',
    'D4EDDA', '155724')

# ============================================================
# 附錄：完整零件圖鑑
# ============================================================
doc.add_page_break()
add_h1(doc, '附錄：本章零件圖鑑')

parts_info = [
    ('608ZZ 滾珠軸承',
     '尺寸：內徑8mm x 外徑22mm x 厚7mm\n'
     '用途：減少轉動摩擦，讓轉盤和太陽爐能輕鬆旋轉\n'
     '外觀：銀色金屬圓環，可以看到裡面的小鋼珠\n'
     '購買：蝦皮搜尋「608ZZ」，一顆約NT$5~10'),
    ('TT 減速馬達',
     '規格：6V供電，減速比1:48，轉速約4RPM\n'
     '用途：提供慢速但力氣大的旋轉動力\n'
     '外觀：黃色塑膠外殼，一端有白色軸\n'
     '購買：蝦皮搜尋「TT 減速馬達」，一顆約NT$35~50'),
    ('小齒輪（馬達用）',
     '規格：適配TT馬達軸徑（約2mm），約10齒\n'
     '用途：裝在馬達軸上，和大齒輪咬合\n'
     '外觀：白色或黃色塑膠小齒輪\n'
     '購買：通常買TT馬達會附，也可單買'),
    ('大齒輪（轉盤用）',
     '規格：直徑6~8cm，約60齒\n'
     '用途：固定在轉盤底部或垂直軸上，被小齒輪驅動\n'
     '外觀：白色或黃色塑膠大齒輪\n'
     '購買：蝦皮搜尋「TT馬達 齒輪組」'),
]

for name, info in parts_info:
    p = doc.add_paragraph()
    r = p.add_run(name)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0, 102, 51)
    doc.add_paragraph(info)
    doc.add_paragraph('')

# ============================================================
# 儲存
# ============================================================
output_path = r'G:\我的雲端硬碟\solar-furnace\tools\科展-太陽爐自動追日系統\第4章-製作轉動載臺-細部說明.docx'
doc.save(output_path)
print('Word 檔已儲存至：' + output_path)
