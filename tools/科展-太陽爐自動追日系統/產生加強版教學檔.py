#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
太陽爐自動追日系統 - 國小科展加強版教學 Word 檔
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# 全域設定
style = doc.styles['Normal']
font = style.font
font.name = '微軟正黑體'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_h1(doc, text):
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 102, 153)

def add_h2(doc, text):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 102, 51)

def add_h3(doc, text):
    h = doc.add_heading(text, level=3)
    for run in h.runs:
        run.font.color.rgb = RGBColor(51, 102, 153)

def add_tip(doc, title, text):
    p = doc.add_paragraph()
    run = p.add_run(title + '：')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 102, 153)
    p.add_run(text)

def add_warn(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(204, 0, 0)

def make_table(doc, headers, rows, hc='1F4E79'):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, hc)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(11)
    for ri, rd in enumerate(rows):
        for ci, val in enumerate(rd):
            table.rows[ri+1].cells[ci].text = val
            for p in table.rows[ri+1].cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return table

# ============================================================
# 封面
# ============================================================
for _ in range(4):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('太陽爐自動追日系統')
r.bold = True
r.font.size = Pt(32)
r.font.color.rgb = RGBColor(0, 102, 153)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('國小科展組裝教學手冊（加強版）')
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = RGBColor(51, 153, 102)

doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('像向日葵一樣，讓太陽爐一直對著太陽！')
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(102, 102, 102)
r.italic = True

doc.add_page_break()

# ============================================================
# 第 1 章
# ============================================================
add_h1(doc, '第 1 章：這是什麼？')

add_h2(doc, '1.1 太陽爐是什麼？')
doc.add_paragraph(
    '太陽爐是一種「不用瓦斯、不用電」的爐子，它用「太陽光」來煮東西！'
    '它的原理就像你用手電筒照放大鏡，光線集中成一個小點，那個點會變得非常熱。'
    '太陽爐就是用一個大大的「凹面鏡」把太陽光集中到一個點（焦點），'
    '讓那裡變得超級熱，熱到可以煮水、煮蛋、甚至爆米花！'
)

add_tip(doc, '小知識',
    '太陽的表面溫度約 5500℃，但地球表面每平方公尺每小時只收到約 1000 瓦的太陽能。'
    '太陽爐的任務就是把分散的太陽能「集中」起來！')

add_h2(doc, '1.2 拋物面集中光線')
dia = doc.add_paragraph()
r = dia.add_run(
    '        太陽光（平行光線）\n'
    '        |    |    |    |    |\n'
    '   \\    |    |    |    |    /\n'
    '     \\  |    |    |    |  /\n'
    '       \\|    |    |    |/\n'
    '         \\  |    |  /\n'
    '           \\|  |/\n'
    '             * 焦點（最熱！100度以上）\n'
    '             鍋子放在這裡'
)
r.font.name = '新細明體'
r.font.size = Pt(11)

add_h2(doc, '1.3 為什麼要追日？')
doc.add_paragraph(
    '太陽在天空中每小時會移動大約 15 度。'
    '如果你的太陽爐固定不動，30 分鐘後太陽就跑掉了，焦點就不在鍋子上了！'
)

dia2 = doc.add_paragraph()
r = dia2.add_run(
    '    10:00太陽    11:00太陽    12:00太陽\n'
    '        *             |             *\n'
    '       /              |              \\\n'
    '      /               |               \\\n'
    '     /                |                \\\n'
    '    /================|================\\\n'
    '    |      太陽爐（如果固定不動）       |\n'
    '    \\================|================/\n'
    '         ^              ^            ^\n'
    '       焦點1          焦點2        焦點3\n'
    '       (對)          (偏了!)       (更偏!)\n\n'
    '    -> 所以我們要讓太陽爐「跟著太陽轉」！'
)
r.font.name = '新細明體'
r.font.size = Pt(10)

doc.add_page_break()

# ============================================================
# 第 2 章
# ============================================================
add_h1(doc, '第 2 章：材料與工具')

add_h2(doc, '2.1 太陽爐本體材料')
make_table(doc,
    ['物品', '規格', '數量', '去哪買', '參考價'],
    [
        ['鏡面壓克力板', '2mm厚 20x20cm', '18片', '光華/蝦皮', '$200~300'],
        ['木合板（做模具）', '1.5cm厚', '1小片', '五金行', '$50'],
        ['銅片', '寬1cm', '36小片', '五金行', '$30'],
        ['螺絲螺帽', '小號', '36組', '五金行', '$30'],
        ['熱熔膠', '膠條', '5支', '文具行', '$30'],
    ])

add_h2(doc, '2.2 追日系統材料')
make_table(doc,
    ['物品', '規格', '數量', '去哪買', '參考價'],
    [
        ['光敏電阻 LDR', 'GL5528', '4顆', '蝦皮', '$5x4'],
        ['LM393 比較器模組', '含旋鈕', '2塊', '蝦皮', '$15x2'],
        ['L293D 馬達驅動', '可控制2馬達', '1塊', '蝦皮', '$45'],
        ['TT 減速馬達', '6V 1:48', '2顆', '蝦皮', '$40x2'],
        ['電阻 10K', '1/4W', '4顆', '電子材料行', '$2x4'],
        ['LED 5mm', '紅/綠 各2', '4顆', '電子材料行', '$3'],
        ['7805 穩壓IC', 'TO-220', '1顆', '電子材料行', '$10'],
        ['電容 100uF', '16V', '2顆', '電子材料行', '$5x2'],
        ['萬用電路板', '5x7cm', '1塊', '電子材料行', '$15'],
        ['6V電池盒', '4xAA附開關', '1個', '蝦皮', '$30'],
        ['AA充電電池', '4顆', '1組', '蝦皮', '$80'],
        ['杜邦線', '公對公+母對母', '各20條', '蝦皮', '$20'],
    ])

add_h2(doc, '2.3 工具清單')
tools = ['線鋸機', '烤箱（可調溫115度）', '熱熔槍', '螺絲起子',
         '三用電表', '電烙鐵', '砂紙', '護目鏡', '隔熱手套', '尺+鉛筆']
for t in tools:
    doc.add_paragraph(t, style='List Bullet')

doc.add_page_break()

# ============================================================
# 第 3 章
# ============================================================
add_h1(doc, '第 3 章：製作太陽爐本體')

add_h2(doc, '3.1 畫拋物線模型')
doc.add_paragraph('打開 GeoGebra 軟體（免費），輸入 y = (1/40) x^2')
doc.add_paragraph('這條線的焦距是 10 公分（從碗底到焦點的距離）')
doc.add_paragraph('把這條線分成 18 等份，每一份就是一片壓克力板的形狀')

add_h2(doc, '3.2 製作模具')
doc.add_paragraph('把列印出來的圖案貼到木合板上')
doc.add_paragraph('用線鋸機沿著線切出來，這就是模具')
add_warn(doc, '警告：線鋸機很鋒利，一定要有大人陪同！')

add_h2(doc, '3.3 裁切壓克力板')
doc.add_paragraph('把鏡面壓克力板放在模具上畫線')
doc.add_paragraph('用線鋸機裁切 18 片')
doc.add_paragraph('用砂紙把邊緣磨光滑')

add_h2(doc, '3.4 加熱彎曲成型')
doc.add_paragraph('烤箱預熱到 115 度')
doc.add_paragraph('把一片壓克力板放到模具上')
doc.add_paragraph('連同模具一起放進烤箱')
doc.add_paragraph('加熱 6.5 分鐘')
doc.add_paragraph('戴隔熱手套取出')
doc.add_paragraph('冷卻 10 分鐘定型')
doc.add_paragraph('重複 18 次')
add_warn(doc, '警告：烤箱很燙！小朋友要請大人幫忙！')

add_h2(doc, '3.5 組合成太陽爐')
doc.add_paragraph('18 片壓克力板排成碗的形狀')
doc.add_paragraph('相鄰兩片用銅片+螺絲鎖在一起')
doc.add_paragraph('碗底用熱熔膠黏合固定')
doc.add_paragraph('檢查縫隙，用熱熔膠補起來')

dia3 = doc.add_paragraph()
r = dia3.add_run(
    '     片1  片2  片3\n'
    '      \\   |   /\n'
    '       \\  |  /\n'
    '    片18-\\|/-片4\n'
    '          * <-- 焦點（碗底中央上方10cm）\n'
    '    片17-/|-\\片5\n'
    '       /  |  \\\n'
    '      /   |   \\\n'
    '    片16  ...  片6\n\n'
    '    每片之間用銅片+螺絲固定，碗底用熱熔膠'
)
r.font.name = '新細明體'
r.font.size = Pt(11)

doc.add_page_break()

# ============================================================
# 第 4 章
# ============================================================
add_h1(doc, '第 4 章：製作轉動載臺')

add_h2(doc, '4.1 水平轉動盤')
doc.add_paragraph('裁切一塊 25x25cm 木合板作為轉盤')
doc.add_paragraph('裁切一塊 30x30cm 木合板作為底座')
doc.add_paragraph('在轉盤四角各放一顆滾珠軸承（608ZZ）')
doc.add_paragraph('滾珠讓轉盤可以輕鬆轉動')

dia4 = doc.add_paragraph()
r = dia4.add_run(
    '  水平轉動盤（上視圖）：\n\n'
    '    +---------------------+\n'
    '    | o               o   |  <-- 4顆滾珠\n'
    '    |                     |\n'
    '    |    +-----------+    |\n'
    '    |    | 太陽爐放這裡 |    |\n'
    '    |    +-----------+    |\n'
    '    |                     |\n'
    '    | o               o   |\n'
    '    +---------------------+\n'
    '           | 底座 |'
)
r.font.name = '新細明體'
r.font.size = Pt(11)

add_h2(doc, '4.2 垂直轉動機構')
doc.add_paragraph('在轉盤兩側豎立兩根木條（高20cm），頂端鑽孔')
doc.add_paragraph('太陽爐兩側裝上短木條，穿過頂端的孔')
doc.add_paragraph('測試：用手輕推太陽爐，它應該可以輕鬆上下轉動')

add_h2(doc, '4.3 安裝馬達')
doc.add_paragraph('水平馬達：固定在底座上，小齒輪對準轉盤的大齒輪')
doc.add_paragraph('垂直馬達：固定在側邊木條上，連接太陽爐的垂直軸')

add_tip(doc, '提醒',
    '馬達要選「減速馬達」（TT馬達 1:48），轉速慢但力氣大。'
    '一般小馬達轉太快，太陽爐會晃來晃去！')

doc.add_page_break()

# ============================================================
# 第 5 章
# ============================================================
add_h1(doc, '第 5 章：追日電路')

add_h2(doc, '5.1 運作原理（像耳朵一樣）')
p = doc.add_paragraph()
r = p.add_run(
    '想像你閉著眼睛，有人從左邊照手電筒。'
    '你的左耳覺得比較亮，右耳比較暗，你就知道光從左邊來，所以你會向左轉。'
)
r.italic = True

doc.add_paragraph(
    '追日系統也是這樣！用兩顆「光敏電阻」（像耳朵）站在左右兩邊，'
    '比較哪邊比較亮，然後決定要向左轉還是向右轉。'
)

add_h2(doc, '5.2 認識零件')

parts = [
    ('光敏電阻（LDR）', '光線感測器。光越強，電阻越小，電壓越高。就像你的皮膚碰一下就知道有沒有光。'),
    ('LM393 比較器', '「比大小」的晶片。比較兩邊電壓，告訴馬達要向左轉還是向右轉。'),
    ('L293D 馬達驅動', '「馬達開關」。比較器說要轉，L293D就打開開關讓馬達轉動，還保護電路不被燒壞。'),
    ('TT 減速馬達', '「慢速但力氣大」的馬達。裡面有齒輪組，把轉速降慢但力氣放大。'),
    ('可變電阻', '可以調整「靈敏度」的旋鈕。太陽偏多少才要轉動，可以調整。'),
]
for name, desc in parts:
    p = doc.add_paragraph()
    r = p.add_run(name + '：')
    r.bold = True
    r.font.color.rgb = RGBColor(0, 51, 153)
    p.add_run(desc)

add_h2(doc, '5.3 電路圖（水平追蹤）')
dia5 = doc.add_paragraph()
r = dia5.add_run(
    '  水平追蹤電路圖：\n\n'
    '        6V 電池\n'
    '           |\n'
    '       +---+---+\n'
    '       | 電源開關|\n'
    '       +---+---+\n'
    '           |\n'
    '     +-----+------------------------+\n'
    '     |     |                         |\n'
    '  +--+--+ ++------+ +----------+    |\n'
    '  | LDR | | 7805  | | LED指示燈 |    |\n'
    '  | 左   | | 穩壓   | | 紅/綠     |    |\n'
    '  +--+--+ +--+-----+ +----+-----+    |\n'
    '     |      | 5V        |            |\n'
    '     |   +--+-----------+--+         |\n'
    '     |   |  LM393 比較器   |         |\n'
    '     |   | (V+ vs V-)      |         |\n'
    '     |   +--------+--------+         |\n'
    '     |            |                   |\n'
    '     |   +--------+--------+         |\n'
    '     |   |  L293D 馬達驅動  |         |\n'
    '     |   +--------+--------+         |\n'
    '     |            |                   |\n'
    '     |     +------+------+           |\n'
    '     |     | 水平減速馬達  |           |\n'
    '     |     |  (東西向)    |           |\n'
    '     |     +-------------+           |\n'
    '     |                                |\n'
    '     +--------------------------------+\n\n'
    '  垂直追蹤電路完全一樣，只是馬達換成控制仰角的那顆！'
)
r.font.name = '新細明體'
r.font.size = Pt(10)

add_h2(doc, '5.4 焊接步驟')
steps = [
    '在萬用電路板上焊上 LM393 比較器',
    '焊接兩顆光敏電阻，中間加一片隔板',
    '焊接可變電阻（旋鈕）',
    '焊接 L293D 馬達驅動晶片',
    '焊接 LED 指示燈',
    '用杜邦線連接馬達',
    '接上電池測試',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'Step {i}：{s}', style='List Number')

add_warn(doc, '警告：電烙鐵很燙（300度以上），一定要有大人陪同！焊接時要戴護目鏡。')

add_h2(doc, '5.5 LDR 安裝位置')
dia6 = doc.add_paragraph()
r = dia6.add_run(
    '  LDR 安裝方式（側面圖）：\n\n'
    '         遮光隔板\n'
    '            |\n'
    '     +------+------+\n'
    '     |      |      |\n'
    '     |  LDR | LDR  |\n'
    '     |  左   | 右   |\n'
    '     |      |      |\n'
    '     +------+------+\n\n'
    '  太陽在正前方：兩邊一樣亮 -> 電壓相等 -> 馬達不動\n'
    '  太陽在左邊：左邊比較亮 -> 電壓不同 -> 馬達向左轉\n'
    '  太陽在右邊：右邊比較亮 -> 電壓不同 -> 馬達向右轉'
)
r.font.name = '新細明體'
r.font.size = Pt(11)

doc.add_page_break()

# ============================================================
# 第 6 章
# ============================================================
add_h1(doc, '第 6 章：組裝全部零件')

steps6 = [
    ('步驟1：固定太陽爐', '把太陽爐固定在轉動載臺的垂直軸上，確保它能上下轉動。'),
    ('步驟2：安裝水平LDR', '把水平追蹤的兩顆光敏電阻安裝在轉盤的東方位置。'),
    ('步驟3：安裝垂直LDR', '把垂直追蹤的兩顆光敏電阻安裝在太陽爐側邊。'),
    ('步驟4：連接電路', '用杜邦線把 LDR -> LM393 -> L293D -> 馬達，一條一條接起來。'),
    ('步驟5：安裝鍋具支架', '在焦點位置（距爐底10cm）裝上鍋具支架。'),
    ('步驟6：接上電源', '裝入4顆AA電池，打開電源開關。'),
    ('步驟7：最終測試', '用手電筒從不同角度照射LDR，觀察馬達是否會轉動。'),
]
for title, desc in steps6:
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0, 102, 51)
    doc.add_paragraph(desc)

dia7 = doc.add_paragraph()
r = dia7.add_run(
    '  組裝完成圖（側面）：\n\n'
    '                 拋物面太陽爐\n'
    '                /    \\\n'
    '               /  *   \\  <-- 焦點（鍋子在這裡）\n'
    '              /        \\\n'
    '             /__________\\\n'
    '                  |\n'
    '         +--------+--------+\n'
    '         |  垂直轉動軸      |  <-- 控制仰角\n'
    '         +--------+--------+\n'
    '                  |\n'
    '    +-------------+-------------+\n'
    '    |             |             |\n'
    '    |    水平轉動盤（東西向）     |  <-- 控制方位角\n'
    '    |    o       o       o     |  <-- 滾珠\n'
    '    +-------------+-------------+\n'
    '    +---------------------------+\n'
    '    |         底座              |\n'
    '    +---------------------------+\n'
    '    +---------------------------+\n'
    '    |  電路板 + 電池 + 馬達      |\n'
    '    +---------------------------+'
)
r.font.name = '新細明體'
r.font.size = Pt(10)

doc.add_page_break()

# ============================================================
# 第 7 章
# ============================================================
add_h1(doc, '第 7 章：測試與調整')

add_h2(doc, '7.1 室內測試')
for s in [
    '用手電筒從不同角度照射LDR',
    '觀察LED指示燈：紅燈亮=向東轉，綠燈亮=向西轉',
    '調整可變電阻（旋鈕），讓馬達不會太靈敏也不會太遲鈍',
    '確認馬達轉動方向正確',
]:
    doc.add_paragraph(s, style='List Number')

add_h2(doc, '7.2 戶外測試')
for s in [
    '選擇晴天，把太陽爐搬到戶外',
    '調整初始角度，讓太陽爐對準太陽',
    '觀察追日系統是否能自動跟隨太陽',
    '每10分鐘記錄一次焦點溫度',
    '記錄2小時，畫出溫度變化圖',
]:
    doc.add_paragraph(s, style='List Number')

add_h2(doc, '7.3 調整技巧')
make_table(doc,
    ['問題', '可能原因', '解決方法'],
    [
        ['馬達一直晃動', '太靈敏了', '把旋鈕轉小一點'],
        ['馬達都不動', '太遲鈍或沒電', '1.檢查電池 2.把旋鈕轉大一點'],
        ['馬達轉反了', '正負極接反', '把馬達的兩條線對調'],
        ['溫度升不高', '焦點沒對準', '調整LDR位置'],
    ])

doc.add_page_break()

# ============================================================
# 第 8 章
# ============================================================
add_h1(doc, '第 8 章：常見問題')

faqs = [
    ('Q：壓克力加熱後怎麼不彎曲？', 'A：溫度不夠！烤箱要到115度，而且要加熱6.5分鐘。不同品牌烤箱有差異，可以多加30秒。'),
    ('Q：18片拼起來有縫隙怎麼辦？', 'A：用熱熔膠補起來就好！小縫隙不影響效果。'),
    ('Q：馬達轉動但太陽爐不動？', 'A：可能是齒輪沒嚙合好，檢查小齒輪有沒有卡到大齒輪。'),
    ('Q：陰天可以測試嗎？', 'A：陰天散射光太多，LDR很難判斷方向，建議等晴天。'),
    ('Q：可以用一般直流馬達嗎？', 'A：不建議！一般馬達轉太快（數千RPM），太陽爐會晃不停。一定要用「減速馬達」。'),
    ('Q：成本大概多少？', 'A：太陽爐本體約NT$250，電路部分約NT$400，總共約NT$650~800。'),
]
for q, a in faqs:
    p = doc.add_paragraph()
    r = p.add_run(q)
    r.bold = True
    r.font.color.rgb = RGBColor(0, 51, 153)
    doc.add_paragraph(a)
    doc.add_paragraph('')

doc.add_page_break()

# ============================================================
# 附錄
# ============================================================
add_h1(doc, '附錄：完整材料採購清單')

make_table(doc,
    ['編號', '物品', '規格', '參考價格', '購買地點'],
    [
        ['1', '鏡面壓克力板', '2mm 20x20cm x18片', 'NT$250', '光華/蝦皮'],
        ['2', 'TT減速馬達', '6V 1:48', 'NT$40x2', '蝦皮'],
        ['3', 'LM393模組', '含旋鈕', 'NT$15x2', '蝦皮'],
        ['4', 'L293D模組', '2馬達驅動', 'NT$45', '蝦皮'],
        ['5', '光敏電阻', 'GL5528', 'NT$5x4', '蝦皮'],
        ['6', '電阻10K', '1/4W', 'NT$2x4', '電子材料行'],
        ['7', 'LED 5mm', '紅/綠各2', 'NT$3', '電子材料行'],
        ['8', '7805穩壓IC', 'TO-220', 'NT$10', '電子材料行'],
        ['9', '電容100uF', '16V', 'NT$5x2', '電子材料行'],
        ['10', '萬用電路板', '5x7cm', 'NT$15', '電子材料行'],
        ['11', '電池盒4xAA', '附開關', 'NT$30', '蝦皮'],
        ['12', 'AA充電電池', '4顆', 'NT$80', '蝦皮'],
        ['13', '杜邦線', '公+母各20', 'NT$20', '蝦皮'],
        ['14', '木合板', '30x30cm', 'NT$50', '五金行'],
        ['15', '銅片+螺絲', '小號', 'NT$30', '五金行'],
        ['16', '608ZZ軸承', '8x22x7mm', 'NT$20x4', '蝦皮'],
    ])

doc.add_paragraph('')
p = doc.add_paragraph()
r = p.add_run('總材料費估算：約 NT$700~1000')
r.bold = True
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0, 102, 51)

# ============================================================
# 儲存
# ============================================================
output_path = r'G:\我的雲端硬碟\solar-furnace\tools\科展-太陽爐自動追日系統\國小科展-太陽爐自動追日系統組裝教學-加強版.docx'
doc.save(output_path)
print('Word 檔已儲存至：' + output_path)
