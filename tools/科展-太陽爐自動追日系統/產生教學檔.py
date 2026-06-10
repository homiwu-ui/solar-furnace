#!/usr/bin/env python3
"""
產生「太陽爐自動追日系統」國小科展教學 Word 檔
小朋友看得懂的一步一步組裝教學
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── 全域設定 ──
style = doc.styles['Normal']
font = style.font
font.name = '微軟正黑體'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

# ============================================================
# 封面
# ============================================================
doc.add_paragraph('')
doc.add_paragraph('')
doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('太陽爐自動追日系統')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0, 102, 153)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('國小科展組裝教學手冊')
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(51, 153, 102)

doc.add_paragraph('')
doc.add_paragraph('')

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('像向日葵一樣，讓太陽爐一直對著太陽！')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(102, 102, 102)

doc.add_page_break()

# ============================================================
# 目錄
# ============================================================
doc.add_heading('目錄', level=1)
toc_items = [
    '第 1 章：這是什麼？— 認識太陽爐',
    '第 2 章：準備工作 — 材料與工具',
    '第 3 章：製作太陽爐本體（拋物面鏡）',
    '第 4 章：製作轉動載臺',
    '第 5 章：製作追日電路',
    '第 6 章：組裝全部零件',
    '第 7 章：測試與調整',
    '第 8 章：常見問題與解決方法',
    '附錄：材料採購清單',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(6)

doc.add_page_break()

# ============================================================
# 第 1 章：這是什麼？
# ============================================================
doc.add_heading('第 1 章：這是什麼？— 認識太陽爐', level=1)

doc.add_heading('1.1 太陽爐是什麼？', level=2)
doc.add_paragraph(
    '太陽爐是一種「不用瓦斯、不用電」的爐子，它用「太陽光」來煮東西！'
)
doc.add_paragraph(
    '它的原理就像你用手電筒照放大鏡，光線集中成一個小點，那個點會變得非常熱。'
    '太陽爐就是用一個大大的「凹面鏡」把太陽光集中到一個點（焦點），讓那裡變得超級熱，'
    '熱到可以煮水、煮蛋、甚至爆米花！'
)

doc.add_heading('1.2 為什麼要「追日」？', level=2)
doc.add_paragraph(
    '太陽在天空中每小時會移動大約 15 度（想像時鐘的時針，一小時轉 30 度的一半）。'
)
doc.add_paragraph(
    '如果你的太陽爐固定不動，30 分鐘後太陽就跑掉了，焦點就不在鍋子上了！'
)
doc.add_paragraph(
    '所以我們要讓太陽爐像「向日葵」一樣，一直對著太陽轉，這就是「自動追日系統」。'
)

# 用文字畫圖
doc.add_heading('太陽移動示意圖', level=3)
diagram = doc.add_paragraph()
diagram.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = diagram.add_run(
    '      ☀️ 太陽\n'
    '       ↗\n'
    '      / 每小時移動 15°\n'
    '     /\n'
    '    /______\n'
    '   📷 太陽爐（如果不動，焦點會跑掉！）'
)
run.font.name = '新細明體'
run.font.size = Pt(11)

doc.add_heading('1.3 我們要做什麼？', level=2)
doc.add_paragraph(
    '我們要做兩件事：'
)
doc.add_paragraph('① 製作一個太陽爐（把太陽光集中）', style='List Number')
doc.add_paragraph('② 製作一個「追日系統」（讓太陽爐自動轉動，一直對著太陽）', style='List Number')

doc.add_page_break()

# ============================================================
# 第 2 章：準備工作
# ============================================================
doc.add_heading('第 2 章：準備工作 — 材料與工具', level=1)

doc.add_heading('2.1 材料清單（太陽爐本體）', level=2)
table = doc.add_table(rows=6, cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['物品', '規格', '數量', '去哪買']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

data = [
    ['鏡面壓克力板', '2mm 厚，約 20×20cm', '18 片', '光華商場 / 蝦皮'],
    ['木合板（做模具）', '1.5cm 厚', '1 小片', '五金行'],
    ['銅片', '寬 1cm', '36 小片', '五金行'],
    ['螺絲螺帽', '小號', '36 組', '五金行'],
    ['熱熔膠', '膠條', '5 支', '文具行 / 蝦皮'],
]
for r, row_data in enumerate(data):
    for c, val in enumerate(row_data):
        table.rows[r+1].cells[c].text = val

doc.add_heading('2.2 材料清單（追日系統）', level=2)
table2 = doc.add_table(rows=12, cols=4)
table2.style = 'Light Grid Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(headers):
    cell = table2.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

data2 = [
    ['光敏電阻（LDR）', '5528 型', '4 顆', '蝦皮 / 電子材料行'],
    ['LM393 比較器模組', '含可變電阻', '2 塊', '蝦皮 / 台灣物聯科技'],
    ['L293D 馬達驅動模組', '可控制 2 顆馬達', '1 塊', '蝦皮 / 台灣物聯科技'],
    ['TT 減速馬達', '6V, 1:48', '2 顆', '蝦皮 / 米羅科技'],
    ['電阻', '10KΩ', '4 顆', '電子材料行'],
    ['LED（紅、綠）', '5mm', '各 2 顆', '電子材料行'],
    ['7805 穩壓 IC', 'TO-220', '1 顆', '電子材料行'],
    ['電容', '100μF 16V', '2 顆', '電子材料行'],
    ['萬用電路板', '5×7cm', '1 塊', '電子材料行'],
    ['6V 電池盒', '4×AA', '1 個', '蝦皮 / 五金行'],
    ['杜邦線', '公對公 + 公對母', '各 20 條', '蝦皮'],
]
for r, row_data in enumerate(data2):
    for c, val in enumerate(row_data):
        table2.rows[r+1].cells[c].text = val

doc.add_heading('2.3 工具清單', level=2)
tools = [
    '線鋸機（或手工線鋸）',
    '烤箱（可調溫，能到 115℃）',
    '熱熔槍',
    '螺絲起子',
    '三用電表',
    '電烙鐵（焊接用）',
    '砂紙',
    '護目鏡',
    '隔熱手套',
    '尺、鉛筆',
]
for t in tools:
    doc.add_paragraph(t, style='List Bullet')

doc.add_page_break()

# ============================================================
# 第 3 章：製作太陽爐本體
# ============================================================
doc.add_heading('第 3 章：製作太陽爐本體（拋物面鏡）', level=1)

doc.add_heading('3.1 什麼是拋物面？', level=2)
doc.add_paragraph(
    '拋物面就像一個「碗」的形狀。當太陽光平行射進來，經過碗面的反射，'
    '所有光線都會集中到一個點，這就是「焦點」。'
)
doc.add_paragraph(
    '我們的太陽爐就是做成這個「碗」的形狀，用 18 片壓克力板拼成一個拋物面。'
)

# 拋物面示意圖
doc.add_heading('拋物面示意圖', level=3)
dia = doc.add_paragraph()
run = dia.add_run(
    '        太陽光（平行光線）\n'
    '        ↓  ↓  ↓  ↓  ↓\n'
    '   ╲    ↓  ↓  ↓  ↓    ╱\n'
    '     ╲  ↓  ↓  ↓  ↓  ╱\n'
    '       ╲↓  ↓  ↓  ↓╱\n'
    '         ╲  ↓  ↓ ╱\n'
    '           ╲↓ ↓╱\n'
    '             ⭐ 焦點（最熱的地方！）\n'
    '             🍳 鍋子放在這裡'
)
run.font.name = '新細明體'
run.font.size = Pt(11)

doc.add_heading('步驟 1：畫出拋物線模型', level=2)
doc.add_paragraph(
    '我們用 GeoGebra（免費電腦軟體）來畫拋物線。'
)
doc.add_paragraph('打開 GeoGebra，輸入：y = (1/40) × x²', style='List Number')
doc.add_paragraph('這條線的「焦距」是 10 公分（從碗底到焦點的距離）', style='List Number')
doc.add_paragraph('把這條線分成 18 等份，每一等份就是一片壓克力板的形狀', style='List Number')
doc.add_paragraph('把每一等份的形狀列印出來，貼到木板上', style='List Number')

doc.add_heading('步驟 2：製作模具', level=2)
doc.add_paragraph(
    '把列印出來的圖案貼到木合板上，用線鋸機沿著線切出來。'
    '這就是我們的「模具」，要用來讓壓克力板彎曲成型。'
)
p = doc.add_paragraph()
run = p.add_run('⚠️ 注意：線鋸機很鋒利，一定要有大人陪同操作！')
run.bold = True
run.font.color.rgb = RGBColor(204, 0, 0)

doc.add_heading('步驟 3：裁切壓克力板', level=2)
doc.add_paragraph('把鏡面壓克力板放在模具上，用鉛筆畫出形狀', style='List Number')
doc.add_paragraph('用線鋸機裁切，總共要切 18 片', style='List Number')
doc.add_paragraph('用砂紙把邊緣磨光滑，避免割傷手', style='List Number')

doc.add_heading('步驟 4：加熱彎曲成型', level=2)
doc.add_paragraph(
    '這是最有趣的步驟！壓克力加熱後會變軟，可以彎曲成型。'
)
doc.add_paragraph('把烤箱預熱到 115℃', style='List Number')
doc.add_paragraph('把一片壓克力板放到模具上', style='List Number')
doc.add_paragraph('連同模具一起放進烤箱', style='List Number')
doc.add_paragraph('加熱 6.5 分鐘', style='List Number')
doc.add_paragraph('戴上隔熱手套取出', style='List Number')
doc.add_paragraph('放在桌上冷卻 10 分鐘，壓克力就定型了！', style='List Number')
doc.add_paragraph('重複 18 次，做完 18 片', style='List Number')

p = doc.add_paragraph()
run = p.add_run('⚠️ 注意：烤箱很熱，一定要戴隔熱手套！小朋友要請大人幫忙！')
run.bold = True
run.font.color.rgb = RGBColor(204, 0, 0)

doc.add_heading('步驟 5：組合成太陽爐', level=2)
doc.add_paragraph('把 18 片壓克力板排成一個碗的形狀', style='List Number')
doc.add_paragraph('相鄰兩片用「銅片 + 螺絲」鎖在一起（上下各一片銅片）', style='List Number')
doc.add_paragraph('碗底（尖端）用熱熔膠黏合固定', style='List Number')
doc.add_paragraph('檢查有沒有縫隙，有的話用熱熔膠補起來', style='List Number')
doc.add_paragraph('最後再撕掉壓克力表面的保護膜', style='List Number')

# 組裝示意圖
doc.add_heading('組裝示意圖', level=3)
dia2 = doc.add_paragraph()
run = dia2.add_run(
    '  18 片壓克力板排成碗狀：\n\n'
    '     片1  片2  片3\n'
    '      ╲   │   ╱\n'
    '       ╲  │  ╱\n'
    '    片18─╲│╱─片4\n'
    '          ⭐ ← 焦點（碗底中央上方 10cm）\n'
    '    片17─╱│╲─片5\n'
    '       ╱  │  ╲\n'
    '      ╱   │   ╲\n'
    '    片16  ...  片6\n\n'
    '  每片之間用銅片+螺絲固定，碗底用熱熔膠'
)
run.font.name = '新細明體'
run.font.size = Pt(11)

doc.add_page_break()

# ============================================================
# 第 4 章：製作轉動載臺
# ============================================================
doc.add_heading('第 4 章：製作轉動載臺', level=1)

doc.add_paragraph(
    '轉動載臺就是太陽爐的「底盤」，它要能讓太陽爐水平轉動（東西向）和垂直轉動（仰角）。'
)

doc.add_heading('4.1 水平轉動盤', level=2)

doc.add_heading('步驟 1：裁切木板', level=3)
doc.add_paragraph('裁切一塊 25cm × 25cm 的木合板作為「轉盤」', style='List Number')
doc.add_paragraph('裁切一塊 30cm × 30cm 的木合板作為「底座」', style='List Number')

doc.add_heading('步驟 2：安裝滾珠', level=3)
doc.add_paragraph(
    '在轉盤的四個角落各放一顆滾珠軸承。'
    '滾珠的作用是讓轉盤可以輕鬆轉動，不會卡住。'
)
doc.add_paragraph('在轉盤四角鑽凹洞，把滾珠嵌進去', style='List Number')
doc.add_paragraph('把轉盤放到大底座上，測試能不能順暢轉動', style='List Number')

# 滾珠示意圖
dia3 = doc.add_paragraph()
run = dia3.add_run(
    '  水平轉動盤（上視圖）：\n\n'
    '    ┌─────────────────────┐\n'
    '    │ ●               ●   │ ← 4 顆滾珠\n'
    '    │                     │\n'
    '    │    ┌───────────┐    │\n'
    '    │    │ 太陽爐放這裡 │    │\n'
    '    │    └───────────┘    │\n'
    '    │                     │\n'
    '    │ ●               ●   │\n'
    '    └─────────────────────┘\n'
    '         ↓ 底座 ↓'
)
run.font.name = '新細明體'
run.font.size = Pt(11)

doc.add_heading('4.2 垂直轉動機構', level=2)
doc.add_paragraph('在轉盤兩側豎立兩根木條（高 20cm），頂端鑽孔', style='List Number')
doc.add_paragraph('太陽爐兩側也裝上短木條，穿過頂端的孔，形成一個「可旋轉的軸」', style='List Number')
doc.add_paragraph('測試：用手輕推太陽爐，它應該可以輕鬆上下轉動', style='List Number')

doc.add_heading('4.3 安裝馬達', level=2)
doc.add_paragraph('水平馬達：固定在底座上，小齒輪對準轉盤的大齒輪', style='List Number')
doc.add_paragraph('垂直馬達：固定在側邊木條上，連接太陽爐的垂直軸', style='List Number')
doc.add_paragraph('用電池測試：馬達轉動時，太陽爐應該會跟著轉', style='List Number')

p = doc.add_paragraph()
run = p.add_run(
    '💡 小提醒：馬達要選「減速馬達」，轉速慢但力氣大。'
    '一般小馬達轉太快，太陽爐會晃來晃去停不下來！'
)
run.font.color.rgb = RGBColor(0, 102, 153)

doc.add_page_break()

# ============================================================
# 第 5 章：製作追日電路
# ============================================================
doc.add_heading('第 5 章：製作追日電路', level=1)

doc.add_paragraph(
    '這是整個專案最「科技」的部分！我們要讓電腦幫太陽爐「看」太陽在哪裡，'
    '然後自動轉動去對準它。'
)

doc.add_heading('5.1 追日電路是怎麼運作的？', level=2)

# 用小朋友看得懂的比喻
p = doc.add_paragraph()
run = p.add_run(
    '想像你閉著眼睛，有人拿手電筒從左邊照你。'
    '你的左耳覺得比較亮，右耳比較暗，你就知道光從左邊來，'
    '所以你會向左轉，讓兩邊一樣亮。'
)
run.italic = True

doc.add_paragraph(
    '追日系統也是這樣！它用兩顆「光敏電阻」（像耳朵）分別站在左右兩邊，'
    '比較哪邊比較亮，然後決定要向左轉還是向右轉。'
)

doc.add_heading('5.2 認識電子零件', level=2)

# 零件說明
parts = [
    ('光敏電阻（LDR）', '一種「光線感測器」。光越強，電阻越小，電壓越高。就像你的皮膚，碰一下就知道有沒有光。'),
    ('LM393 比較器', '一個「比大小」的晶片。它會比較兩邊的電壓，告訴馬達「向左轉」還是「向右轉」。'),
    ('L293D 馬達驅動', '一個「馬達開關」。比較器說要轉，L293D 就打開開關讓馬達轉動。它還會保護電路不被燒壞。'),
    ('TT 減速馬達', '一種「慢速但力氣大」的馬達。裡面有齒輪組，把轉速降慢但力氣放大。'),
    ('可變電阻（旋鈕）', '可以調整「靈敏度」的旋鈕。轉一下，就可以決定太陽偏多少才要轉動。'),
    ('7805 穩壓 IC', '把 6V 電池的電壓穩定在 5V，保護比較器不被高壓燒壞。'),
]

for name, desc in parts:
    p = doc.add_paragraph()
    run = p.add_run(f'【{name}】')
    run.bold = True
    run.font.size = Pt(12)
    p.add_run(f'\n{desc}')

doc.add_heading('5.3 電路圖（水平追蹤部分）', level=2)

# 文字電路圖
dia4 = doc.add_paragraph()
run = dia4.add_run(
    '  水平追蹤電路（小朋友看得懂版）：\n\n'
    '        6V 電池\n'
    '           │\n'
    '       ┌───┴───┐\n'
    '       │ 電源開關│\n'
    '       └───┬───┘\n'
    '           │\n'
    '     ┌─────┼────────────────────────┐\n'
    '     │     │                         │\n'
    '  ┌──┴──┐ ┌┴─────┐ ┌──────────┐    │\n'
    '  │ LDR │ │ 7805 │ │  LED指示燈 │    │\n'
    '  │ 左   │ │ 穩壓  │ │  (紅/綠)  │    │\n'
    '  └──┬──┘ └──┬───┘ └────┬─────┘    │\n'
    '     │       │ 5V       │           │\n'
    '     │    ┌──┴──────────┴──┐        │\n'
    '     │    │   LM393 比較器  │        │\n'
    '     │    │   (V+ vs V-)    │        │\n'
    '     │    └────────┬────────┘        │\n'
    '     │             │                  │\n'
    '     │    ┌────────▼────────┐        │\n'
    '     │    │   L293D 馬達驅動  │        │\n'
    '     │    └────────┬────────┘        │\n'
    '     │             │                  │\n'
    '     │      ┌──────▼──────┐          │\n'
    '     │      │  水平減速馬達  │          │\n'
    '     │      │   (東西向)    │          │\n'
    '     │      └─────────────┘          │\n'
    '     │                                │\n'
    '     └────────────────────────────────┘\n\n'
    '  垂直追蹤電路完全一樣，只是馬達換成控制仰角的那顆！'
)
run.font.name = '新細明體'
run.font.size = Pt(10)

doc.add_heading('5.4 焊接步驟', level=2)

doc.add_paragraph('在萬用電路板上焊上 LM393 比較器', style='List Number')
doc.add_paragraph('焊接兩顆光敏電阻，中間加一片隔板（讓兩邊光線分開）', style='List Number')
doc.add_paragraph('焊接可變電阻（旋鈕），用來調整靈敏度', style='List Number')
doc.add_paragraph('焊接 L293D 馬達驅動晶片', style='List Number')
doc.add_paragraph('焊接 LED 指示燈（紅色=向東轉，綠色=向西轉）', style='List Number')
doc.add_paragraph('用杜邦線連接馬達', style='List Number')
doc.add_paragraph('接上電池測試', style='List Number')

p = doc.add_paragraph()
run = p.add_run('⚠️ 焊接注意事項：')
run.bold = True
run.font.color.rgb = RGBColor(204, 0, 0)
doc.add_paragraph('電烙鐵很燙（300℃以上），一定要有大人陪同！', style='List Bullet')
doc.add_paragraph('焊接時要戴護目鏡，避免焊錫飛濺', style='List Bullet')
doc.add_paragraph('焊接完要用三用電表確認沒有短路', style='List Bullet')

doc.add_heading('5.5 LDR 感測器安裝位置', level=2)

# LDR 安裝示意圖
dia5 = doc.add_paragraph()
run = dia5.add_run(
    '  LDR 安裝方式（側面圖）：\n\n'
    '         遮光隔板\n'
    '            │\n'
    '     ┌──────┼──────┐\n'
    '     │      │      │\n'
    '     │  LDR │ LDR  │\n'
    '     │  左   │ 右   │\n'
    '     │      │      │\n'
    '     └──────┴──────┘\n\n'
    '  當太陽在正前方：兩邊一樣亮 → 電壓相等 → 馬達不動\n'
    '  當太陽在左邊：左邊比較亮 → 電壓不同 → 馬達向左轉\n'
    '  當太陽在右邊：右邊比較亮 → 電壓不同 → 馬達向右轉'
)
run.font.name = '新細明體'
run.font.size = Pt(11)

doc.add_page_break()

# ============================================================
# 第 6 章：組裝全部零件
# ============================================================
doc.add_heading('第 6 章：組裝全部零件', level=1)

doc.add_paragraph('現在，我們要把前面做好的三個部分組合在一起！')

doc.add_heading('組裝順序', level=2)

steps = [
    ('步驟 1：固定太陽爐', '把組裝好的拋物面太陽爐固定在轉動載臺的垂直軸上。確保它能上下轉動。'),
    ('步驟 2：安裝水平 LDR', '把水平追蹤的兩顆光敏電阻安裝在轉盤的東方位置（因為太陽從東邊升起，這樣比較早能感測到）。'),
    ('步驟 3：安裝垂直 LDR', '把垂直追蹤的兩顆光敏電阻安裝在太陽爐側邊，負責感測仰角。'),
    ('步驟 4：連接電路', '用杜邦線把 LDR → LM393 → L293D → 馬達，一條一條接起來。'),
    ('步驟 5：安裝鍋具支架', '在焦點位置（距爐底 10cm）裝上鍋具支架，把黑色鍋子放上去。'),
    ('步驟 6：接上電源', '裝入 4 顆 AA 電池，打開電源開關。'),
    ('步驟 7：最終測試', '用手電筒從不同角度照射 LDR，觀察馬達是否會轉動。'),
]

for title, desc in steps:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}')
    run.bold = True
    run.font.size = Pt(12)
    p.add_run(f'\n{desc}')

# 組裝完成圖
doc.add_heading('組裝完成示意圖', level=3)
dia6 = doc.add_paragraph()
run = dia6.add_run(
    '  整體組裝完成圖（側面）：\n\n'
    '                 📷 拋物面太陽爐\n'
    '                ╱    ╲\n'
    '               ╱  ⭐  ╲  ← 焦點（鍋子在這裡）\n'
    '              ╱        ╲\n'
    '             ╱__________╲\n'
    '                  │\n'
    '         ┌────────┴────────┐\n'
    '         │  垂直轉動軸      │ ← 控制仰角\n'
    '         └────────┬────────┘\n'
    '                  │\n'
    '    ┌─────────────┼─────────────┐\n'
    '    │             │             │\n'
    '    │    水平轉動盤（東西向）     │ ← 控制方位角\n'
    '    │    ●       ●       ●     │ ← 滾珠\n'
    '    └─────────────┴─────────────┘\n'
    '    ┌─────────────────────────┐\n'
    '    │         底座             │\n'
    '    └─────────────────────────┘\n'
    '                  │\n'
    '    ┌─────────────┴─────────────┐\n'
    '    │  電路板 + 電池 + 馬達       │\n'
    '    └─────────────────────────┘'
)
run.font.name = '新細明體'
run.font.size = Pt(10)

doc.add_page_break()

# ============================================================
# 第 7 章：測試與調整
# ============================================================
doc.add_heading('第 7 章：測試與調整', level=1)

doc.add_heading('7.1 室內測試', level=2)
doc.add_paragraph('用手電筒從不同角度照射 LDR', style='List Number')
doc.add_paragraph('觀察 LED 指示燈：紅燈亮=向東轉，綠燈亮=向西轉', style='List Number')
doc.add_paragraph('調整可變電阻（旋鈕），讓馬達不會太靈敏（一直晃動）也不會太遲鈍', style='List Number')
doc.add_paragraph('確認馬達轉動方向正確', style='List Number')

doc.add_heading('7.2 戶外測試', level=2)
doc.add_paragraph('選擇晴天，把太陽爐搬到戶外', style='List Number')
doc.add_paragraph('調整初始角度，讓太陽爐對準太陽', style='List Number')
doc.add_paragraph('觀察追日系統是否能自動跟隨太陽', style='List Number')
doc.add_paragraph('每 10 分鐘記錄一次焦點溫度', style='List Number')
doc.add_paragraph('記錄 2 小時，畫出溫度變化圖', style='List Number')

doc.add_heading('7.3 調整技巧', level=2)
table3 = doc.add_table(rows=5, cols=3)
table3.style = 'Light Grid Accent 1'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
headers3 = ['問題', '可能原因', '解決方法']
for i, h in enumerate(headers3):
    cell = table3.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

adjust_data = [
    ['馬達一直晃動不停', '太靈敏了', '把可變電阻（旋鈕）轉小一點'],
    ['馬達都不動', '太遲鈍或沒電', '1. 檢查電池  2. 把旋鈕轉大一點'],
    ['馬達轉反了', '正負極接反了', '把馬達的兩條線對調'],
    ['溫度升不高', '焦點沒對準', '調整 LDR 位置，確認焦點在鍋子上'],
]
for r, row_data in enumerate(adjust_data):
    for c, val in enumerate(row_data):
        table3.rows[r+1].cells[c].text = val

doc.add_page_break()

# ============================================================
# 第 8 章：常見問題
# ============================================================
doc.add_heading('第 8 章：常見問題與解決方法', level=1)

faqs = [
    ('Q：壓克力加熱後怎麼不彎曲？', 'A：溫度不夠！烤箱要到 115℃，而且要加熱 6.5 分鐘。不同品牌烤箱溫度有差異，可以試著多加 30 秒。'),
    ('Q：18 片拼起來有縫隙怎麼辦？', 'A：用熱熔膠補起來就好！小縫隙不影響效果。'),
    ('Q：馬達轉動但太陽爐不動？', 'A：可能是齒輪沒嚙合好，檢查小齒輪有沒有卡到大齒輪。'),
    ('Q：陰天可以測試嗎？', 'A：陰天散射光太多，LDR 很難判斷方向，建議等晴天再測。'),
    ('Q：可以用一般直流馬達嗎？', 'A：不建議！一般馬達轉太快（數千RPM），太陽爐會晃不停。一定要用「減速馬達」。'),
    ('Q：成本大概多少？', 'A：太陽爐本體約 NT$ 200~300（壓克力板），電路部分約 NT$ 350~500，總共約 NT$ 600~800。'),
]

for q, a in faqs:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 51, 153)
    doc.add_paragraph(a)
    doc.add_paragraph('')

doc.add_page_break()

# ============================================================
# 附錄：材料採購清單
# ============================================================
doc.add_heading('附錄：材料採購清單（含參考價格）', level=1)

table4 = doc.add_table(rows=16, cols=5)
table4.style = 'Light Grid Accent 1'
table4.alignment = WD_TABLE_ALIGNMENT.CENTER
headers4 = ['編號', '物品', '規格', '參考價格', '購買地點']
for i, h in enumerate(headers4):
    cell = table4.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

purchase_data = [
    ['1', '鏡面壓克力板 2mm', '20×20cm × 18片', 'NT$ 200~300', '光華/蝦皮'],
    ['2', 'TT 減速馬達 6V', '1:48', 'NT$ 40 × 2', '蝦皮'],
    ['3', 'LM393 比較器模組', '含旋鈕', 'NT$ 15 × 2', '蝦皮/台灣物聯科技'],
    ['4', 'L293D 馬達驅動模組', '可控制2馬達', 'NT$ 45', '蝦皮'],
    ['5', '光敏電阻 LDR 5528', '5~10mm', 'NT$ 5 × 4', '蝦皮'],
    ['6', '電阻 10KΩ', '¼W', 'NT$ 2 × 4', '電子材料行'],
    ['7', 'LED 5mm 紅/綠', '各2顆', 'NT$ 3', '電子材料行'],
    ['8', '7805 穩壓 IC', 'TO-220', 'NT$ 10', '電子材料行'],
    ['9', '電容 100μF 16V', '電解電容', 'NT$ 5 × 2', '電子材料行'],
    ['10', '萬用電路板', '5×7cm', 'NT$ 15', '電子材料行'],
    ['11', '6V 電池盒 4×AA', '附開關', 'NT$ 30', '蝦皮'],
    ['12', 'AA 充電電池', '4顆', 'NT$ 80', '蝦皮/超商'],
    ['13', '杜邦線', '公對公+公對母', 'NT$ 20', '蝦皮'],
    ['14', '木合板', '30×30cm', 'NT$ 50', '五金行'],
    ['15', '銅片+螺絲', '小號', 'NT$ 30', '五金行'],
]
for r, row_data in enumerate(purchase_data):
    for c, val in enumerate(row_data):
        table4.rows[r+1].cells[c].text = val

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('總材料費估算：約 NT$ 700~1000（不含太陽能板）')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 102, 51)

# ============================================================
# 儲存
# ============================================================
output_path = r'G:\我的雲端硬碟\solar-furnace\tools\科展-太陽爐自動追日系統\國小科展-太陽爐自動追日系統組裝教學.docx'
doc.save(output_path)
print(f'Word 檔已儲存至：{output_path}')
